"""Post-processing du memoire.docx généré par pandoc — normes Pr MBOG ENSPD."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("memoire_complet.docx")

# --- Marges (Pr MBOG : Haut 3, Gauche 3, Bas 3, Droite 2.5) ---
for section in doc.sections:
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# --- Style Normal (corps) ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- Headings ---
heading_cfg = {
    'Heading 1': (Pt(16), True,  True,  RGBColor(0,0,0)),   # CHAPITRE — Majuscule, gras
    'Heading 2': (Pt(14), True,  False, RGBColor(0,0,0)),   # Section — Gras
    'Heading 3': (Pt(12), True,  False, RGBColor(0x33,0x33,0x33)),  # Sous-section — Gras
    'Heading 4': (Pt(12), False, False, RGBColor(0x33,0x33,0x33)),  # Paragraphe — Normal
}

for name, (size, bold, caps, color) in heading_cfg.items():
    if name in doc.styles:
        s = doc.styles[name]
        s.font.name  = 'Times New Roman'
        s.font.size  = size
        s.font.bold  = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after  = Pt(8)
        s.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

# --- Tables ---
for table in doc.tables:
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True

# --- S'assurer que tous les runs sont en Times New Roman ---
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.name is None:
            run.font.name = 'Times New Roman'

doc.save("memoire_complet.docx")
print("Mémoire formaté : Times New Roman 12pt, interligne 1.15, marges ENSPD, tableaux OK.")
