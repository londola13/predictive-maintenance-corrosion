"""Post-processing du .docx généré par pandoc pour formatage académique."""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document("revue_litterature.docx")

# --- Marges ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# --- Style par défaut (corps) ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# --- Heading styles ---
heading_sizes = {
    'Heading 1': (Pt(16), True, RGBColor(0, 0, 0)),
    'Heading 2': (Pt(14), True, RGBColor(0, 0, 0)),
    'Heading 3': (Pt(13), True, RGBColor(0x33, 0x33, 0x33)),
    'Heading 4': (Pt(12), True, RGBColor(0x33, 0x33, 0x33)),
}

for name, (size, bold, color) in heading_sizes.items():
    if name in doc.styles:
        s = doc.styles[name]
        s.font.name = 'Times New Roman'
        s.font.size = size
        s.font.bold = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(8)

# --- Table formatting ---
for table in doc.tables:
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    # Bold first row (header)
    if table.rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

# --- All runs: ensure Times New Roman ---
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.name is None:
            run.font.name = 'Times New Roman'

doc.save("revue_litterature.docx")
print("Formatting applied: Times New Roman 12pt, 1.5 spacing, academic margins.")
