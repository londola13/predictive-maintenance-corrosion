"""
Générateur Word — Revue de littérature M2
Maintenance prédictive de la corrosion par XGBoost
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy

OUTPUT = "D:/Claude code/predictive-maintenance-corrosion/memoire/revue_litterature.docx"

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=6, line=276):
    """Espacement paragraphe (line = 276 = simple, 360 = 1.5)"""
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_heading(doc, text, level):
    para = doc.add_heading(text, level=level)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.runs[0]
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return para

def add_body(doc, text):
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    set_paragraph_spacing(para, before=0, after=6, line=360)  # interligne 1.5
    return para

def add_omml_equation(doc, omml_str):
    """Insère une équation Word (OMML) centrée"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(para, before=6, after=6)
    oMath = etree.fromstring(omml_str)
    para._p.append(oMath)
    return para

def add_reaction(doc, text):
    """Réaction chimique — police monospace centrée"""
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    run.font.italic = True
    set_paragraph_spacing(para, before=4, after=4)
    return para

def add_legend(doc, items: list):
    """Légende d'équation — liste de définitions"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("Avec :")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    set_paragraph_spacing(para, before=2, after=2)
    for symbol, definition in items:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p.add_run(f"{symbol}")
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r1.font.bold = True
        r2 = p.add_run(f"  =  {definition}")
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
        set_paragraph_spacing(p, before=0, after=2)

def add_note(doc, text):
    """Note encadrée (contexte lien prototype)"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(f"📌  {text}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x6D, 0x3A)
    set_paragraph_spacing(para, before=4, after=4)
    return para

# ─────────────────────────────────────────────────────────────
# ÉQUATIONS OMML
# ─────────────────────────────────────────────────────────────

# Namespace OMML
NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'

# Loi de Faraday : m = (M · I · t) / (n · F)
OMML_FARADAY = f"""<m:oMath {NS}>
  <m:r><m:rPr><m:sty m:val="p"/></m:rPr><m:t>m&#160;=&#160;</m:t></m:r>
  <m:f>
    <m:num><m:r><m:t>M&#160;&#xB7;&#160;I&#160;&#xB7;&#160;t</m:t></m:r></m:num>
    <m:den><m:r><m:t>n&#160;&#xB7;&#160;F</m:t></m:r></m:den>
  </m:f>
</m:oMath>"""

# Taux de corrosion : CR = (87,6 · Δm) / (ρ · A · t)
OMML_CR = f"""<m:oMath {NS}>
  <m:r><m:rPr><m:sty m:val="p"/></m:rPr><m:t>CR&#160;(mm/an)&#160;=&#160;</m:t></m:r>
  <m:f>
    <m:num><m:r><m:t>87,6&#160;&#xB7;&#160;&#x394;m</m:t></m:r></m:num>
    <m:den><m:r><m:t>&#x3C1;&#160;&#xB7;&#160;A&#160;&#xB7;&#160;t</m:t></m:r></m:den>
  </m:f>
</m:oMath>"""

# Butler-Volmer
OMML_BV = f"""<m:oMath {NS}>
  <m:r><m:t>i&#160;=&#160;i</m:t></m:r>
  <m:sSub>
    <m:e><m:r><m:t></m:t></m:r></m:e>
    <m:sub><m:r><m:t>corr</m:t></m:r></m:sub>
  </m:sSub>
  <m:r><m:t>&#160;&#xB7;&#160;</m:t></m:r>
  <m:d>
    <m:dPr><m:begChr m:val="["/><m:endChr m:val="]"/></m:dPr>
    <m:e>
      <m:sSup>
        <m:e><m:r><m:t>e</m:t></m:r></m:e>
        <m:sup>
          <m:f>
            <m:num><m:r><m:t>&#x3B1;</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>a</m:t></m:r></m:sub></m:sSub><m:r><m:t>&#xB7;F&#xB7;&#x3B7;</m:t></m:r></m:num>
            <m:den><m:r><m:t>R&#xB7;T</m:t></m:r></m:den>
          </m:f>
        </m:sup>
      </m:sSup>
      <m:r><m:t>&#160;&#x2212;&#160;</m:t></m:r>
      <m:sSup>
        <m:e><m:r><m:t>e</m:t></m:r></m:e>
        <m:sup>
          <m:f>
            <m:num><m:r><m:t>&#x2212;&#x3B1;</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>c</m:t></m:r></m:sub></m:sSub><m:r><m:t>&#xB7;F&#xB7;&#x3B7;</m:t></m:r></m:num>
            <m:den><m:r><m:t>R&#xB7;T</m:t></m:r></m:den>
          </m:f>
        </m:sup>
      </m:sSup>
    </m:e>
  </m:d>
</m:oMath>"""

# ─────────────────────────────────────────────────────────────
# DOCUMENT
# ─────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── PAGE DE TITRE ──────────────────────────────────────────
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("REVUE DE LITTÉRATURE")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(
        "Système de maintenance prédictive de la corrosion par apprentissage automatique :\n"
        "conception d'une sonde ER low-cost, acquisition IoT\n"
        "et prédiction du taux de corrosion par XGBoost"
    )
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(13)
    r2.font.italic = True

    doc.add_paragraph()
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run("Mémoire de Master 2 — ENSPD Douala\nMaintenance Industrielle — Corrosion/Érosion Oil & Gas")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

    doc.add_page_break()

    # ── PARTIE 1 ───────────────────────────────────────────────
    add_heading(doc, "PARTIE 1 — LA CORROSION : PHÉNOMÈNE, ENJEUX ET RÉPONSES INDUSTRIELLES", 1)

    # 1.1
    add_heading(doc, "1.1  Mécanismes électrochimiques fondamentaux", 2)

    add_body(doc,
        "La corrosion des métaux est un phénomène naturel et irréversible par lequel un matériau "
        "métallique tend à retourner à son état thermodynamiquement stable, c'est-à-dire à l'état "
        "oxydé sous lequel il se trouve naturellement dans la croûte terrestre. Le fer, par exemple, "
        "existe à l'état naturel sous forme d'hématite (Fe₂O₃) ou de magnétite (Fe₃O₄). La "
        "transformation de ce minerai en acier par métallurgie constitue un apport d'énergie que la "
        "corrosion vient progressivement dissiper."
    )

    add_body(doc,
        "La définition normative de référence est celle de la norme ISO 8044 : « la corrosion est une "
        "interaction physico-chimique entre un métal et son environnement qui entraîne des modifications "
        "des propriétés du métal et qui peut conduire à une dégradation significative de la fonction du "
        "métal, de l'environnement ou du système technique dont ils font partie. »"
    )

    # 1.1.1
    add_heading(doc, "1.1.1  Nature électrochimique de la corrosion", 3)

    add_body(doc,
        "Dans un milieu électrolytique (solution aqueuse, sol humide, condensat), la corrosion procède "
        "par un mécanisme de pile galvanique dans lequel deux réactions électrochimiques couplées se "
        "produisent simultanément à la surface du métal."
    )

    add_body(doc, "Réaction anodique — oxydation (dissolution du métal) :")
    add_reaction(doc, "Fe  →  Fe²⁺  +  2e⁻")

    add_body(doc,
        "Le métal cède des électrons et des ions métalliques passent en solution. "
        "C'est la zone anodique qui se corrode."
    )

    add_body(doc, "Réaction cathodique — réduction (consommation des électrons) :")
    add_body(doc, "En milieu acide (pH < 4), réduction des protons :")
    add_reaction(doc, "2H⁺  +  2e⁻  →  H₂ ↑")
    add_body(doc, "En milieu neutre ou légèrement acide, réduction de l'oxygène dissous :")
    add_reaction(doc, "O₂  +  2H₂O  +  4e⁻  →  4OH⁻")

    add_body(doc,
        "Ces deux réactions ne peuvent pas exister l'une sans l'autre : les électrons produits à "
        "l'anode doivent obligatoirement être consommés à la cathode. C'est ce couplage qui définit "
        "la densité de courant de corrosion (i_corr), paramètre central de toute quantification."
    )

    add_note(doc,
        "Dans le prototype de laboratoire, l'environnement acide (vinaigre + NaCl, pH ≈ 2,5–3) "
        "active simultanément ces deux réactions sur le fil de fer de la sonde ER."
    )

    # 1.1.2
    add_heading(doc, "1.1.2  Loi de Faraday — lien entre courant et perte de matière", 3)

    add_body(doc,
        "La loi de Faraday (1834) établit la relation quantitative entre le courant électrique échangé "
        "et la masse de métal dissous. Elle constitue le fondement théorique de toutes les méthodes de "
        "mesure électrochimiques présentées au Chapitre 2 :"
    )

    add_omml_equation(doc, OMML_FARADAY)

    add_legend(doc, [
        ("m", "masse de métal dissous (g)"),
        ("M", "masse molaire du métal  —  fer : 55,85 g/mol"),
        ("I", "courant de corrosion (A)"),
        ("t", "durée d'exposition (s)"),
        ("n", "nombre d'électrons échangés  —  2 pour Fe → Fe²⁺"),
        ("F", "constante de Faraday = 96 485 C/mol"),
    ])

    add_body(doc,
        "Le taux de corrosion (CR — Corrosion Rate), exprimé en mm/an conformément à la norme "
        "ASTM G1, est calculé à partir de la perte de masse par unité de surface et de temps :"
    )

    add_omml_equation(doc, OMML_CR)

    add_legend(doc, [
        ("Δm", "perte de masse (mg)"),
        ("ρ",  "densité du métal (g/cm³)  —  acier : 7,87 g/cm³"),
        ("A",  "aire exposée (cm²)"),
        ("t",  "durée d'exposition (heures)"),
        ("87,6", "facteur de conversion dimensionnelle (ASTM G1)"),
    ])

    add_note(doc,
        "Cette formule est la référence industrielle universelle. Elle sera utilisée dans ce travail "
        "pour convertir les variations de résistance électrique mesurées par la sonde ER en taux de "
        "corrosion exprimés en mm/an — variable cible du modèle XGBoost."
    )

    # 1.1.3
    add_heading(doc, "1.1.3  Potentiel de corrosion et diagrammes de Pourbaix", 3)

    add_body(doc,
        "Tout métal immergé dans un électrolyte adopte spontanément un potentiel mixte (E_corr), dit "
        "potentiel de corrosion, résultant de l'équilibre cinétique entre la dissolution anodique et la "
        "réduction cathodique. Les diagrammes de Pourbaix (potentiel E en fonction du pH) délimitent "
        "les domaines de stabilité thermodynamique d'un métal dans un environnement aqueux en trois zones :"
    )

    items = [
        ("Zone de corrosion",  "les ions métalliques (Fe²⁺, Fe³⁺) sont stables en solution — le métal se dissout"),
        ("Zone de passivité",  "un film d'oxyde protecteur (Fe₂O₃) se forme et ralentit la corrosion"),
        ("Zone d'immunité",   "le métal est thermodynamiquement stable — aucune corrosion"),
    ]
    for label, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{label} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=3)

    add_note(doc,
        "Dans le prototype, l'environnement acide (pH ≈ 2,5–3) place le fil de fer dans la zone de "
        "corrosion active du diagramme de Pourbaix — ce qui garantit une dissolution mesurable sur "
        "la durée de l'expérience (28 jours)."
    )

    # 1.1.4
    add_heading(doc, "1.1.4  Cinétique de corrosion — équation de Butler-Volmer", 3)

    add_body(doc,
        "La vitesse réelle de corrosion ne dépend pas seulement de la thermodynamique (diagramme de "
        "Pourbaix) mais aussi de la cinétique électrochimique, décrite par l'équation de Butler-Volmer :"
    )

    add_omml_equation(doc, OMML_BV)

    add_legend(doc, [
        ("i",           "densité de courant (A/m²)"),
        ("i_corr",      "densité de courant de corrosion (A/m²)"),
        ("αa, αc",      "coefficients de transfert anodique et cathodique"),
        ("F",           "constante de Faraday = 96 485 C/mol"),
        ("η",           "surtension = E − E_corr (V)"),
        ("R",           "constante des gaz = 8,314 J/mol·K"),
        ("T",           "température absolue (K)"),
    ])

    add_body(doc,
        "En pratique industrielle, cette équation est exploitée par la méthode de polarisation "
        "linéaire (LPR) qui permet de mesurer i_corr en temps réel — méthode décrite au Chapitre 2. "
        "Elle justifie également l'influence de la température comme facteur aggravant, "
        "paramètre directement intégré dans le modèle XGBoost au Chapitre 4."
    )

    # ── 1.2 ────────────────────────────────────────────────────
    add_heading(doc, "1.2  Classification des formes de corrosion", 2)

    add_body(doc,
        "La corrosion ne se manifeste pas de manière uniforme. Elle prend des formes variées selon "
        "la nature du métal, la composition de l'environnement et les conditions mécaniques. La "
        "classification internationale de référence est établie par l'AMPP (ex-NACE International) "
        "et reprise dans la norme NACE SP0775. Cette classification est essentielle pour tout "
        "ingénieur en maintenance : chaque forme de corrosion impose une stratégie de surveillance "
        "et une méthode de mesure spécifiques."
    )

    # Tableau des formes de corrosion
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Forme", "Mécanisme", "Localisation typique", "Norme / Référence"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr[i].paragraphs[0].runs[0].font.size = Pt(11)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    rows_data = [
        ("Généralisée (uniforme)",
         "Dissolution uniforme sur toute la surface exposée",
         "Pipelines acier au carbone, réservoirs",
         "ASTM G1, NACE SP0775"),
        ("Par piqûres",
         "Attaque localisée créant des cavités profondes ; film passif percé localement par les ions Cl⁻",
         "Acier inoxydable en milieu chloruré",
         "ASTM G46, ISO 11463"),
        ("Galvanique",
         "Couplage électrochimique entre deux métaux de potentiels différents dans un électrolyte",
         "Raccords bimétalliques, structures offshore",
         "ASTM G82, NACE TM0107"),
        ("Sous contrainte (SCC)",
         "Combinaison d'une contrainte mécanique de traction et d'un environnement corrosif",
         "Tubes sous pression, soudures",
         "NACE TM0177, ISO 7539"),
        ("Érosion-corrosion",
         "Attaque accélérée par l'écoulement du fluide qui détruit le film protecteur",
         "Coudes, vannes, pompes",
         "ASTM G76, NACE 35100"),
        ("CO₂ (sweet corrosion)",
         "CO₂ dissous forme H₂CO₃ (acide carbonique) → attaque acide du fer",
         "Pipelines Oil & Gas, têtes de puits",
         "De Waard-Milliams (1975)"),
        ("H₂S (sour corrosion)",
         "H₂S provoque la fragilisation par hydrogène (SSC) et la fissuration (HIC)",
         "Puits à gaz acide, raffineries",
         "NACE MR0175 / ISO 15156"),
        ("MIC (corrosion microbienne)",
         "Activité de bactéries (SRB, IRB) modifiant les conditions électrochimiques locales",
         "Fonds de réservoirs, conduites enterrées",
         "NACE TM0212"),
    ]

    for row_data in rows_data:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    add_body(doc,
        "Dans le secteur Oil & Gas, les deux formes les plus prévalentes sont la corrosion CO₂ "
        "(sweet corrosion) et la corrosion H₂S (sour corrosion). La corrosion CO₂ est régie par "
        "le modèle semi-empirique de De Waard et Milliams (1975, révisé 1991), qui reste la "
        "référence industrielle pour l'estimation du taux de corrosion en fonction de la pression "
        "partielle de CO₂ et de la température. Ce modèle sera discuté en détail à la section 1.3."
    )

    add_body(doc,
        "La corrosion H₂S, régie par la norme NACE MR0175 / ISO 15156, introduit un mécanisme "
        "supplémentaire de fragilisation par hydrogène : les atomes H issus de la réaction "
        "cathodique pénètrent dans le réseau cristallin de l'acier et provoquent des fissures "
        "internes (HIC — Hydrogen Induced Cracking) ou de surface (SSC — Sulfide Stress Cracking). "
        "Ces phénomènes ne sont pas mesurables par la sonde ER et nécessitent des contrôles "
        "non-destructifs spécifiques (UT, TOFD)."
    )

    add_body(doc,
        "Concernant la MIC (Microbiologically Influenced Corrosion), bien que significative dans "
        "les pipelines industriels — notamment dans les fonds de réservoirs où les bactéries "
        "sulfato-réductrices (SRB) prospèrent en conditions anaérobies — elle ne sera pas "
        "reproduite dans ce prototype. L'environnement de laboratoire (vinaigre + NaCl) ne "
        "présente pas les conditions microbiologiques requises. Cette forme de corrosion est "
        "mentionnée pour exhaustivité du panorama mais reste hors du périmètre du système "
        "de mesure développé."
    )

    add_note(doc,
        "Le prototype reproduit une corrosion généralisée en milieu acide (CH₃COOH + NaCl, "
        "pH ≈ 2,5–3), mécanistiquement analogue à la corrosion CO₂ par sa composante acide "
        "(réduction des protons H⁺). Cette analogie sera justifiée quantitativement à la "
        "section 3.3 lors de la conception de l'environnement corrosif contrôlé."
    )

    # ── 1.3 ────────────────────────────────────────────────────
    add_heading(doc, "1.3  Facteurs influents sur le taux de corrosion", 2)

    add_body(doc,
        "Le taux de corrosion n'est pas une constante intrinsèque du matériau : il résulte de "
        "l'interaction dynamique entre le métal et son environnement. La maîtrise de ces facteurs "
        "est fondamentale à deux titres : d'une part, elle guide la conception de l'environnement "
        "corrosif contrôlé du prototype ; d'autre part, ces mêmes facteurs constituent les variables "
        "d'entrée (features) du modèle XGBoost développé au Chapitre 4."
    )

    # 1.3.1 Température
    add_heading(doc, "1.3.1  Température", 3)

    add_body(doc,
        "La température est le facteur cinétique dominant. Son influence sur la vitesse de corrosion "
        "suit une loi d'Arrhenius, qui traduit l'accélération des réactions chimiques avec la "
        "chaleur :"
    )

    OMML_ARRHENIUS = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>CR&#160;&#x221D;&#160;</m:t></m:r>
      <m:sSup>
        <m:e><m:r><m:t>e</m:t></m:r></m:e>
        <m:sup>
          <m:f>
            <m:num><m:r><m:t>&#x2212;Ea</m:t></m:r></m:num>
            <m:den><m:r><m:t>R&#160;&#xB7;&#160;T</m:t></m:r></m:den>
          </m:f>
        </m:sup>
      </m:sSup>
    </m:oMath>"""
    add_omml_equation(doc, OMML_ARRHENIUS)

    add_legend(doc, [
        ("Ea", "énergie d'activation (J/mol)"),
        ("R",  "constante des gaz = 8,314 J/mol·K"),
        ("T",  "température absolue (K)"),
    ])

    add_body(doc,
        "En pratique industrielle, une règle empirique largement utilisée stipule qu'une augmentation "
        "de 10°C double approximativement la vitesse de corrosion. Cette règle est formalisée dans "
        "le modèle de De Waard-Milliams (1975, révisé 1991) pour la corrosion CO₂ :"
    )

    OMML_DEWAARD = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>log(CR)&#160;=&#160;5,8&#160;&#x2212;&#160;</m:t></m:r>
      <m:f>
        <m:num><m:r><m:t>1710</m:t></m:r></m:num>
        <m:den><m:r><m:t>T</m:t></m:r></m:den>
      </m:f>
      <m:r><m:t>&#160;+&#160;0,67&#160;&#xB7;&#160;log(pCO&#x2082;)</m:t></m:r>
    </m:oMath>"""
    add_omml_equation(doc, OMML_DEWAARD)

    add_legend(doc, [
        ("CR",   "taux de corrosion (mm/an)"),
        ("T",    "température absolue (K)"),
        ("pCO₂", "pression partielle de CO₂ (bar)"),
    ])

    add_body(doc,
        "Ce modèle, bien que semi-empirique et limité aux conditions de pipeline Oil & Gas, reste "
        "la référence industrielle pour une première estimation rapide du taux de corrosion. "
        "Dans le prototype, la température est maintenue entre 30°C et 55°C à l'aide d'une "
        "résistance chauffante d'aquarium, et mesurée en continu par le capteur DS18B20."
    )

    # 1.3.2 pH
    add_heading(doc, "1.3.2  pH et acidité", 3)

    add_body(doc,
        "Le pH est le paramètre qui quantifie la concentration en ions H⁺ dans la solution. "
        "En milieu acide (pH < 7), la réaction cathodique de réduction des protons est favorisée, "
        "accélérant la dissolution anodique du métal. En dessous de pH 4, la corrosion du fer "
        "devient très agressive car le film d'oxyde protecteur (Fe₂O₃) se dissout. "
        "Au-dessus de pH 10, une couche passive stable se forme, inhibant naturellement la corrosion."
    )

    add_body(doc,
        "Dans le contexte Oil & Gas, le pH est directement lié à la teneur en CO₂ et H₂S dissous. "
        "Le CO₂ forme de l'acide carbonique (H₂CO₃, pKa₁ = 6,35) et l'H₂S forme de l'acide "
        "sulfhydrique (H₂S, pKa₁ = 7,02), tous deux abaissant le pH de l'eau de formation. "
        "Dans notre prototype, le vinaigre blanc (acide acétique 6–14%) maintient un pH ≈ 2,5–3, "
        "mesuré par bandelette pH ou sonde analogique connectée à l'ADS1115."
    )

    # 1.3.3 Pression partielle CO2
    add_heading(doc, "1.3.3  Pression partielle de CO₂ (pCO₂)", 3)

    add_body(doc,
        "Le CO₂ dissous dans l'eau de formation constitue l'agent corrosif principal des pipelines "
        "de production pétrolière. Il réagit avec l'eau selon :"
    )
    add_reaction(doc, "CO₂  +  H₂O  →  H₂CO₃  →  H⁺  +  HCO₃⁻")

    add_body(doc,
        "La pression partielle de CO₂ (pCO₂) est définie par la loi de Henry : "
        "pCO₂ = x_CO₂ × P_totale, où x_CO₂ est la fraction molaire de CO₂ dans le gaz. "
        "Selon le modèle de De Waard-Milliams, le taux de corrosion est proportionnel à "
        "pCO₂^0,67. En pratique industrielle, une pCO₂ supérieure à 0,2 bar est considérée "
        "comme corrosive ; au-delà de 0,5 bar, la corrosion est sévère (NACE SP0192)."
    )

    add_body(doc,
        "Dans le prototype de laboratoire, la pCO₂ n'est pas mesurée directement — un capteur "
        "de CO₂ dissous dépasserait largement le budget disponible. Elle est contrôlée "
        "indirectement via le pH, conformément à la relation de Henderson-Hasselbalch, et "
        "simulée par l'addition de bicarbonate de soude en réaction avec le vinaigre."
    )

    # 1.3.4 Teneur en eau BSW
    add_heading(doc, "1.3.4  Teneur en eau (BSW)", 3)

    add_body(doc,
        "Le BSW (Basic Sediment and Water) désigne la fraction volumique d'eau dans le fluide "
        "transporté. L'eau est l'électrolyte indispensable à la corrosion : sans eau libre au "
        "contact du métal, les réactions électrochimiques ne peuvent pas se produire. "
        "Un BSW élevé augmente la conductivité ionique de la phase aqueuse et accélère les "
        "échanges de charges à l'interface métal/électrolyte."
    )

    add_body(doc,
        "En dessous de 30% BSW, la phase continue est généralement huileuse et le métal est "
        "naturellement protégé. Au-delà de 60–70% BSW, la phase aqueuse devient continue "
        "et la corrosion s'accélère significativement. Ce paramètre est mesuré en continu "
        "dans les installations industrielles par des analyseurs BSW dédiés (capteur "
        "capacitif ou micro-ondes)."
    )

    # 1.3.5 Vitesse d'écoulement
    add_heading(doc, "1.3.5  Vitesse d'écoulement et érosion-corrosion", 3)

    add_body(doc,
        "La vitesse du fluide influence la corrosion selon deux mécanismes antagonistes. "
        "D'un côté, un débit élevé favorise l'apport d'oxygène et renouvelle l'électrolyte "
        "au contact du métal, accélérant la corrosion. De l'autre, en dessous d'une vitesse "
        "minimale, les dépôts (boues, cire, sédiments) s'accumulent et créent des piles de "
        "concentration locales encore plus agressives."
    )

    add_body(doc,
        "Au-delà d'une vitesse critique (typiquement 3–5 m/s pour l'acier au carbone), "
        "le film protecteur d'oxyde est mécaniquement arraché par le fluide — c'est "
        "l'érosion-corrosion. Ce phénomène est particulièrement sévère aux coudes, "
        "réductions de section et vannes partiellement ouvertes."
    )

    # 1.3.6 Inhibiteurs
    add_heading(doc, "1.3.6  Inhibiteurs de corrosion", 3)

    add_body(doc,
        "Les inhibiteurs de corrosion sont des substances chimiques qui, ajoutées en faible "
        "concentration dans l'environnement corrosif, réduisent significativement le taux "
        "de corrosion. Leur efficacité est quantifiée par le taux d'inhibition η :"
    )

    OMML_ETA = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>&#x3B7;&#160;(%)&#160;=&#160;</m:t></m:r>
      <m:f>
        <m:num><m:r><m:t>CR&#x209B;&#x2090;&#x2099;&#x209B;&#160;&#x2212;&#160;CR&#x1D43;&#x1D9C;&#x1D49;&#x1D9C;</m:t></m:r></m:num>
        <m:den><m:r><m:t>CR&#x209B;&#x2090;&#x2099;&#x209B;</m:t></m:r></m:den>
      </m:f>
      <m:r><m:t>&#160;&#xD7;&#160;100</m:t></m:r>
    </m:oMath>"""
    add_omml_equation(doc, OMML_ETA)

    add_legend(doc, [
        ("CR_sans", "taux de corrosion sans inhibiteur (mm/an)"),
        ("CR_avec", "taux de corrosion avec inhibiteur (mm/an)"),
        ("η",       "taux d'inhibition (%) — objectif industriel : η > 90%"),
    ])

    add_body(doc,
        "Les inhibiteurs se classifient selon leur mécanisme d'action en trois catégories "
        "(ASTM G170, NACE TM0374) :"
    )

    inh_data = [
        ("Anodiques",
         "Forment un film passif sur la zone anodique (chromates, nitrites). "
         "Très efficaces mais toxiques — usage restreint."),
        ("Cathodiques",
         "Précipitent sur les zones cathodiques et limitent la réduction de l'O₂ "
         "(sels de zinc, polyphosphates)."),
        ("Mixtes (films adsorbés)",
         "S'adsorbent sur toute la surface métallique et bloquent simultanément "
         "les zones anodiques et cathodiques. Classe dominante en industrie pétrolière "
         "(amines, imidazolines, inhibiteurs organiques naturels)."),
    ]

    for nom, desc in inh_data:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{nom} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=3)

    add_body(doc,
        "Les inhibiteurs organiques naturels constituent un axe de recherche actif, "
        "particulièrement pertinent dans le contexte africain où des ressources végétales "
        "locales (extraits d'Azadirachta indica, d'Hibiscus sabdariffa, huiles végétales) "
        "ont démontré des propriétés inhibitrices significatives (η > 80%) dans plusieurs "
        "études publiées (Okafor et al., 2010 ; Umoren et al., 2013). Dans ce travail, "
        "l'huile végétale est utilisée comme inhibiteur de phase 2 du protocole expérimental."
    )

    add_note(doc,
        "La dose d'inhibiteur recommandée constitue la SORTIE principale du modèle XGBoost. "
        "Le système prédit le taux de corrosion futur et recommande automatiquement la dose "
        "optimale (0, 0,5, 1 ou 2 mL/L) pour maintenir η > 90% — boucle fermée "
        "mesure → prédiction → action de protection."
    )

    # 1.3.7 Tableau récapitulatif
    add_heading(doc, "1.3.7  Synthèse — Facteurs et features XGBoost", 3)

    add_body(doc,
        "Le tableau suivant synthétise les facteurs influents, leur effet sur le taux de "
        "corrosion et leur statut dans le modèle de prédiction :"
    )

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    for i, txt in enumerate(["Facteur", "Effet sur CR", "Mesuré dans prototype", "Feature XGBoost"]):
        hdr2[i].text = txt
        hdr2[i].paragraphs[0].runs[0].font.bold = True
        hdr2[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        hdr2[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    synth_data = [
        ("Température (T)",       "CR × 2 par +10°C",              "DS18B20 (continu)",            "Oui — temp"),
        ("pH / acidité",          "CR ↑ quand pH ↓",               "Bandelette ou sonde pH",       "Oui — ph"),
        ("pCO₂",                  "CR ∝ pCO₂^0,67",                "Indirect via pH",              "Proxy — ph"),
        ("BSW (teneur en eau)",   "CR ↑ quand BSW ↑",              "Solution 100% aqueuse",        "Fixe = 100%"),
        ("Vitesse écoulement",    "CR ↑ (érosion-corrosion)",      "Non — bac statique",           "Non"),
        ("Inhibiteur",            "CR ↓ si η > 90%",               "Dose ajoutée manuellement",    "Oui — dose_inhibiteur"),
        ("Résistance sonde ER",   "R ↑ quand CR ↑ (fil s'amince)", "ADS1115 + pont Wheatstone",    "Oui — resistance, delta_R"),
    ]

    for row_data in synth_data:
        row = table2.add_row()
        for i, txt in enumerate(row_data):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 1.4 ────────────────────────────────────────────────────
    add_heading(doc, "1.4  Coûts et enjeux mondiaux de la corrosion", 2)

    add_body(doc,
        "La corrosion représente l'un des défis économiques les plus coûteux de l'industrie "
        "mondiale. L'étude de référence NACE IMPACT (International Measures of Prevention, "
        "Application, and Economics of Corrosion Technologies), publiée en 2016, évalue le coût "
        "annuel mondial de la corrosion à 2 500 milliards de dollars USD, soit environ 3,4% du "
        "PIB mondial. Ce chiffre dépasse le coût combiné des catastrophes naturelles annuelles "
        "et illustre l'ampleur du problème."
    )

    add_body(doc,
        "Dans le secteur Oil & Gas spécifiquement, la corrosion représente entre 25% et 35% "
        "de l'ensemble des défaillances de pipelines. Aux États-Unis, la base de données PHMSA "
        "(Pipeline and Hazardous Materials Safety Administration) recense annuellement entre 200 "
        "et 400 incidents significatifs directement attribués à la corrosion sur les réseaux de "
        "pipelines. Le coût d'un incident majeur dépasse fréquemment 10 millions USD, incluant "
        "les travaux de réparation, la dépollution environnementale et les pertes de production."
    )

    add_body(doc,
        "L'étude NACE IMPACT souligne également que 25 à 30% de ces coûts seraient évitables "
        "par l'application systématique des meilleures pratiques de gestion de la corrosion, "
        "parmi lesquelles la surveillance en temps réel et la maintenance prédictive occupent "
        "une place centrale. C'est précisément dans cette perspective que s'inscrit le présent "
        "travail : démontrer qu'une approche low-cost basée sur l'IoT et le machine learning "
        "peut contribuer significativement à la réduction de ces coûts, y compris dans des "
        "contextes à ressources limitées tels que l'Afrique subsaharienne."
    )

    add_body(doc,
        "Dans le contexte camerounais et africain, l'enjeu est particulièrement critique. "
        "Les infrastructures pétrolières (COTCO, Perenco, Total Energies Cameroun) opèrent "
        "dans des environnements tropicaux humides qui accélèrent la corrosion, avec des "
        "températures ambiantes élevées (28–35°C) et une saison des pluies intense favorisant "
        "la corrosion atmosphérique et galvanique. L'accès aux systèmes de surveillance "
        "industriels certifiés étant coûteux, le développement de solutions locales "
        "adaptées constitue un enjeu stratégique de souveraineté technique."
    )

    # ── 1.5 ────────────────────────────────────────────────────
    add_heading(doc, "1.5  Stratégies de protection contre la corrosion", 2)

    add_body(doc,
        "Face aux mécanismes de corrosion décrits précédemment, l'industrie a développé "
        "un ensemble de stratégies de protection complémentaires. La norme NACE SP0169 "
        "et les recommandations AMPP classifient ces stratégies en cinq grandes catégories, "
        "présentées ici dans une logique allant du plus passif au plus actif."
    )

    add_heading(doc, "1.5.1  Vue d'ensemble des approches", 3)

    table_prot = doc.add_table(rows=1, cols=3)
    table_prot.style = 'Table Grid'
    for i, txt in enumerate(["Stratégie", "Principe", "Limitation principale"]):
        table_prot.rows[0].cells[i].text = txt
        table_prot.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_prot.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_prot.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    prot_data = [
        ("Sélection des matériaux", "Utiliser des alliages résistants (inox, Inconel, duplex)", "Coût très élevé"),
        ("Revêtements et peintures", "Barrière physique entre métal et environnement", "Passif, vieillissement, discontinuités"),
        ("Protection cathodique (CP)", "Forcer le métal en zone d'immunité électrochimique", "Infrastructure électrique requise"),
        ("Inhibiteurs de corrosion", "Film chimique adsorbé bloquant les réactions", "Dosage continu, coût chimiques"),
        ("Contrôle de l'environnement", "Dégazage O₂, déshumidification, pH control", "Applicable uniquement en circuit fermé"),
    ]
    for row_d in prot_data:
        row = table_prot.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "1.5.2  Revêtements et peintures anti-corrosion", 3)

    add_body(doc,
        "Les revêtements constituent la première ligne de défense dans la majorité des "
        "infrastructures industrielles. Ils agissent comme une barrière physique empêchant "
        "le contact direct entre le métal et l'environnement corrosif. Les principaux types "
        "utilisés dans l'industrie pétrolière sont les revêtements époxy (application liquide, "
        "norme ISO 12944), le FBE (Fusion Bonded Epoxy, norme ASTM A1046) pour les pipelines "
        "enterrés, et la métallisation par projection thermique de zinc ou d'aluminium "
        "(norme ISO 2063) pour les structures offshore."
    )

    add_body(doc,
        "La limite fondamentale des revêtements est leur caractère passif et non adaptatif : "
        "une fois posé, le revêtement ne s'ajuste pas aux variations des conditions opératoires. "
        "Toute discontinuité (holiday, défaut d'application, vieillissement) crée un point "
        "d'attaque corrosive concentrée, souvent plus sévère que sans revêtement (effet "
        "d'anode de petite surface). C'est pourquoi les revêtements sont systématiquement "
        "combinés avec la protection cathodique sur les pipelines entérrés."
    )

    add_body(doc,
        "Cette stratégie étant purement passive et ne générant aucune donnée mesurable "
        "par notre système, elle ne sera pas développée davantage dans ce mémoire. "
        "Elle constitue néanmoins un prérequis de la conception des pipelines que tout "
        "ingénieur en corrosion doit maîtriser."
    )

    add_heading(doc, "1.5.3  Protection cathodique (CP)", 3)

    add_body(doc,
        "La protection cathodique est l'une des méthodes les plus efficaces et les plus "
        "répandues pour protéger les structures métalliques enterrées ou immergées. Son "
        "principe repose sur la thermodynamique électrochimique : en abaissant le potentiel "
        "du métal à protéger jusqu'à la zone d'immunité du diagramme de Pourbaix, toute "
        "dissolution anodique devient thermodynamiquement impossible."
    )

    add_body(doc,
        "Deux technologies coexistent en industrie (NACE SP0169) :"
    )

    cp_data = [
        ("Anode sacrificielle",
         "Un métal moins noble (zinc, magnésium, aluminium) est connecté électriquement "
         "à la structure à protéger. Ce métal, plus anodique, se corrode préférentiellement "
         "en protégeant la structure. Avantages : aucune alimentation électrique requise, "
         "installation simple. Limitation : durée de vie limitée de l'anode, efficacité "
         "décroissante. Norme EN 13174, NACE SP0387."),
        ("Courant imposé (ICCP)",
         "Un courant électrique continu est imposé de l'extérieur via un redresseur "
         "et des anodes inertes (platine, MMO). Le potentiel de protection visé est "
         "−850 mV/CSE (Electrode au Sulfate de Cuivre) pour l'acier au carbone selon "
         "NACE SP0169. Avantages : puissance ajustable, très longue durée de vie. "
         "Limitation : infrastructure électrique requise, risque de surprotection "
         "(fragilisation par hydrogène)."),
    ]
    for nom, desc in cp_data:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{nom} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=4)

    add_note(doc,
        "Dans le prototype, la Phase 3 du protocole expérimental reproduit le principe "
        "de l'anode sacrificielle : une plaque de zinc est connectée au fil de fer "
        "(acier), créant un couple galvanique où le zinc se corrode préférentiellement. "
        "Le modèle XGBoost enregistre la variable protection_cathodique (booléen) "
        "et quantifie la réduction de CR associée."
    )

    add_heading(doc, "1.5.4  Inhibiteurs de corrosion — Application industrielle", 3)

    add_body(doc,
        "La section 1.3.6 a présenté les mécanismes d'action et la classification des "
        "inhibiteurs. Cette section se concentre sur leur application industrielle et "
        "les enjeux de dosage, directement reliés à la sortie du modèle XGBoost."
    )

    add_body(doc,
        "En industrie pétrolière, les inhibiteurs sont injectés en continu dans le fluide "
        "transporté via des pompes doseuses calibrées. La concentration typique est de "
        "l'ordre de 10 à 100 ppm (mg/L). Un sous-dosage entraîne une protection insuffisante "
        "(η < 70%), tandis qu'un surdosage est économiquement pénalisant et peut interférer "
        "avec les traitements en aval (désémulsification, traitement des eaux). La norme "
        "NACE TM0374 définit les protocoles d'évaluation de l'efficacité des inhibiteurs "
        "en laboratoire par tests rotatifs (wheel test)."
    )

    add_body(doc,
        "Le défi industriel n'est donc pas seulement de choisir le bon inhibiteur, mais "
        "d'ajuster sa dose en temps réel en fonction des conditions opératoires variables "
        "(température, débit, BSW). C'est précisément ce problème d'optimisation dynamique "
        "que le modèle XGBoost de ce travail vise à résoudre : prédire le taux de corrosion "
        "futur pour recommander la dose d'inhibiteur minimale garantissant η > 90%."
    )

    # ── 1.6 ────────────────────────────────────────────────────
    add_heading(doc, "1.6  Cadre normatif applicable", 2)

    add_body(doc,
        "La mesure, la surveillance et la gestion de la corrosion sont encadrées par un "
        "ensemble de normes internationales édictées par l'ASTM International, l'AMPP "
        "(ex-NACE), l'ISO, l'API et l'ASME. Le tableau suivant présente les normes "
        "directement applicables à ce travail :"
    )

    table_norm = doc.add_table(rows=1, cols=3)
    table_norm.style = 'Table Grid'
    for i, txt in enumerate(["Norme", "Organisme", "Objet"]):
        table_norm.rows[0].cells[i].text = txt
        table_norm.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_norm.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_norm.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    norm_data = [
        ("ASTM G1",           "ASTM",  "Nettoyage et préparation des éprouvettes de corrosion"),
        ("ASTM G31",          "ASTM",  "Essais d'immersion en laboratoire"),
        ("ASTM G46",          "ASTM",  "Inspection et caractérisation de la corrosion par piqûres"),
        ("ASTM G96-90",       "ASTM",  "Surveillance en ligne par méthodes électriques et électrochimiques (ER, LPR)"),
        ("ASTM G170",         "ASTM",  "Évaluation et qualification des inhibiteurs de corrosion"),
        ("NACE SP0775",       "AMPP",  "Préparation et utilisation des coupons de corrosion en industrie"),
        ("NACE TM0169",       "AMPP",  "Évaluation de la corrosion par tests en laboratoire (coupons)"),
        ("NACE TM0374",       "AMPP",  "Évaluation des inhibiteurs par test rotatif (wheel test)"),
        ("NACE MR0175/ISO 15156", "AMPP/ISO", "Matériaux pour environnements H₂S (sour service)"),
        ("NACE SP0169",       "AMPP",  "Protection cathodique des pipelines enterrés"),
        ("ISO 8044",          "ISO",   "Vocabulaire de la corrosion"),
        ("ISO 8407",          "ISO",   "Procédures de nettoyage des éprouvettes corrodées"),
        ("API 570",           "API",   "Inspection des systèmes de tuyauterie en service"),
        ("API 580/581",       "API",   "Inspection basée sur le risque (RBI) — méthodologie"),
        ("ASME B31.3",        "ASME",  "Tuyauterie de process — calcul durée de vie résiduelle"),
    ]
    for row_d in norm_data:
        row = table_norm.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Conclusion Partie 1
    add_heading(doc, "Conclusion de la Partie 1", 2)
    add_body(doc,
        "Cette première partie a établi les fondements scientifiques et industriels indispensables "
        "à la compréhension du système développé dans ce travail. La corrosion est un phénomène "
        "électrochimique quantifiable (loi de Faraday, taux de corrosion en mm/an), multiforme "
        "(8 types identifiés par la NACE), influencé par des paramètres mesurables (T, pH, pCO₂, "
        "dose inhibiteur) et encadré par un corpus normatif international rigoureux."
    )
    add_body(doc,
        "Les stratégies de protection disponibles — revêtements, protection cathodique, inhibiteurs "
        "— sont efficaces mais réactives : elles ne s'adaptent pas dynamiquement aux variations des "
        "conditions opératoires. Le défi central que ce travail adresse est précisément ce passage "
        "du réactif au prédictif : anticiper l'évolution du taux de corrosion pour recommander "
        "automatiquement la dose d'inhibiteur optimale au bon moment. La partie suivante présente "
        "les méthodes de mesure qui rendent cette prédiction possible."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PARTIE 2
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PARTIE 2 — MESURE ET SURVEILLANCE DE LA CORROSION EN INDUSTRIE", 1)

    add_heading(doc, "2.1  Classification des méthodes de mesure (ASTM G96)", 2)

    add_body(doc,
        "La norme ASTM G96-90 (Standard Guide for Online Monitoring of Corrosion in Plant "
        "Equipment) constitue le référentiel international de classification des méthodes "
        "de surveillance de la corrosion. Elle distingue deux axes orthogonaux : le mode "
        "d'acquisition (hors-ligne vs en ligne) et le caractère de la mesure "
        "(direct vs indirect, destructif vs non-destructif)."
    )

    table_meth = doc.add_table(rows=1, cols=4)
    table_meth.style = 'Table Grid'
    for i, txt in enumerate(["Méthode", "Mode", "Type", "Grandeur mesurée"]):
        table_meth.rows[0].cells[i].text = txt
        table_meth.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_meth.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_meth.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    meth_data = [
        ("Gravimétrie (coupons)",      "Hors-ligne", "Destructif / Direct",       "Perte de masse (mg)"),
        ("LPR",                        "En ligne",   "Non-destructif / Indirect",  "Résistance de polarisation Rp (Ω)"),
        ("EIS",                        "En ligne",   "Non-destructif / Indirect",  "Impédance complexe Z(ω)"),
        ("Sonde ER",                   "En ligne",   "Destructif / Direct",        "Résistance électrique R (Ω)"),
        ("UT épaisseur",               "Hors-ligne", "Non-destructif / Direct",    "Épaisseur de paroi (mm)"),
        ("PAUT (Phased Array)",        "Hors-ligne", "Non-destructif / Direct",    "Image 2D/3D de l'épaisseur"),
        ("Guided Wave (GWT)",          "En ligne",   "Non-destructif / Direct",    "Épaisseur moyenne sur longue portée"),
    ]
    for row_d in meth_data:
        row = table_meth.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "2.2  Méthodes hors-ligne — Gravimétrie et coupons", 2)

    add_body(doc,
        "La méthode gravimétrique, décrite par les normes ASTM G1 et NACE SP0775, est la "
        "technique de référence historique pour la mesure du taux de corrosion. Elle consiste "
        "à exposer des éprouvettes métalliques (coupons) de masse et de surface connues dans "
        "le fluide corrosif pendant une durée déterminée, puis à mesurer la perte de masse "
        "après nettoyage selon le protocole ISO 8407."
    )

    add_body(doc,
        "Le taux de corrosion est calculé directement par la formule ASTM G1 présentée à la "
        "section 1.1.2. Cette méthode présente l'avantage d'une grande simplicité de mise en "
        "œuvre et d'un coût minimal. Elle est cependant fondamentalement discontinue : les "
        "coupons doivent être retirés pour être pesés, ce qui ne permet pas un suivi en temps "
        "réel de l'évolution de la corrosion. Les intervalles de changement de coupons sont "
        "typiquement de 30 à 90 jours en industrie, ce qui masque les variations rapides liées "
        "aux changements de conditions opératoires."
    )

    add_heading(doc, "2.3  Méthodes électrochimiques en ligne", 2)

    add_heading(doc, "2.3.1  Polarisation linéaire (LPR)", 3)

    add_body(doc,
        "La polarisation linéaire (LPR — Linear Polarization Resistance) est une méthode "
        "électrochimique permettant de mesurer le taux de corrosion instantané en temps réel, "
        "sans consommation significative du métal (mesure non-destructive). Elle repose sur "
        "la relation de Stern-Geary (1957) :"
    )

    OMML_STERN = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>i</m:t></m:r>
      <m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>corr</m:t></m:r></m:sub></m:sSub>
      <m:r><m:t>&#160;=&#160;</m:t></m:r>
      <m:f>
        <m:num><m:r><m:t>B</m:t></m:r></m:num>
        <m:den><m:r><m:t>R</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>p</m:t></m:r></m:sub></m:sSub></m:den>
      </m:f>
    </m:oMath>"""
    add_omml_equation(doc, OMML_STERN)

    add_legend(doc, [
        ("i_corr", "densité de courant de corrosion (A/m²)"),
        ("B",      "constante de Stern-Geary (V) — typiquement 0,026 V pour acier en milieu acide"),
        ("Rp",     "résistance de polarisation (Ω·m²) — mesurée expérimentalement"),
    ])

    add_body(doc,
        "En pratique, Rp est mesurée en appliquant une petite perturbation de potentiel "
        "(±10–20 mV par rapport à E_corr) et en mesurant le courant résultant. La méthode "
        "LPR est très rapide (une mesure en 1–2 minutes) et adaptée à la surveillance "
        "continue. Sa principale limitation est qu'elle ne fonctionne qu'en milieu "
        "électrolytique conducteur (eau saline, condensat) et donne des résultats peu "
        "fiables en présence d'inhibiteurs adsorbés (modification de B)."
    )

    add_heading(doc, "2.3.2  Spectroscopie d'impédance électrochimique (EIS)", 3)

    add_body(doc,
        "La spectroscopie d'impédance électrochimique (EIS — Electrochemical Impedance "
        "Spectroscopy) est la méthode électrochimique la plus riche en information. Elle "
        "consiste à appliquer une perturbation sinusoïdale de faible amplitude (5–10 mV) "
        "sur une large gamme de fréquences (10⁻² à 10⁵ Hz) et à mesurer la réponse en "
        "courant. Le rapport tension/courant à chaque fréquence définit l'impédance "
        "complexe Z(ω), représentée dans le diagramme de Nyquist ou de Bode."
    )

    add_body(doc,
        "L'ajustement d'un circuit électrique équivalent (circuit de Randles) aux données "
        "expérimentales permet d'extraire la résistance de transfert de charge Rct (liée à "
        "i_corr via Stern-Geary), la résistance de la solution Rs, et la capacité de double "
        "couche Cdl (indicateur de l'état de surface). L'EIS est la méthode de référence "
        "pour l'évaluation des inhibiteurs en laboratoire (ASTM G170) car elle caractérise "
        "complètement l'interface métal/électrolyte. Sa complexité et son coût d'équipement "
        "la réservent cependant à la recherche et aux laboratoires spécialisés."
    )

    add_heading(doc, "2.4  Sondes à résistance électrique (ER) — Méthode centrale", 2)

    add_heading(doc, "2.4.1  Principe physique", 3)

    add_body(doc,
        "La méthode de surveillance par résistance électrique (ER — Electrical Resistance) "
        "est définie par la norme ASTM G96-90 comme une technique de surveillance en ligne "
        "basée sur la mesure de la résistance électrique d'un élément sensible métallique "
        "exposé au milieu corrosif. Son principe repose sur la loi de Pouillet :"
    )

    OMML_POUILLET = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>R&#160;=&#160;</m:t></m:r>
      <m:f>
        <m:num><m:r><m:t>&#x3C1;&#160;&#xB7;&#160;L</m:t></m:r></m:num>
        <m:den><m:r><m:t>A</m:t></m:r></m:den>
      </m:f>
    </m:oMath>"""
    add_omml_equation(doc, OMML_POUILLET)

    add_legend(doc, [
        ("R", "résistance électrique de l'élément sensible (Ω)"),
        ("ρ", "résistivité électrique du métal (Ω·m) — fer : 1,0 × 10⁻⁷ Ω·m"),
        ("L", "longueur de l'élément sensible (m)"),
        ("A", "section transversale de l'élément (m²)"),
    ])

    add_body(doc,
        "Lorsque l'élément sensible corrode, sa section A diminue (perte de matière) tandis "
        "que sa longueur L et sa résistivité ρ restent constantes. Il en résulte une "
        "augmentation de la résistance R directement proportionnelle à la perte de matière. "
        "Pour un élément filaire de rayon r(t) décroissant avec le temps :"
    )

    OMML_ER_DERIVE = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>r(t)&#160;=&#160;</m:t></m:r>
      <m:rad>
        <m:radPr><m:degHide m:val="1"/></m:radPr>
        <m:deg/>
        <m:e>
          <m:f>
            <m:num><m:r><m:t>&#x3C1;&#160;&#xB7;&#160;L</m:t></m:r></m:num>
            <m:den><m:r><m:t>&#x3C0;&#160;&#xB7;&#160;R(t)</m:t></m:r></m:den>
          </m:f>
        </m:e>
      </m:rad>
    </m:oMath>"""
    add_omml_equation(doc, OMML_ER_DERIVE)

    add_body(doc,
        "Le taux de corrosion est alors calculé comme la vitesse de diminution du rayon "
        "de l'élément sensible, convertie en mm/an par la relation :"
    )

    OMML_CR_ER = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>CR&#160;(mm/an)&#160;=&#160;</m:t></m:r>
      <m:f>
        <m:num><m:r><m:t>r&#x2080;&#160;&#x2212;&#160;r(t)</m:t></m:r></m:num>
        <m:den><m:r><m:t>t</m:t></m:r></m:den>
      </m:f>
      <m:r><m:t>&#160;&#xD7;&#160;8760</m:t></m:r>
    </m:oMath>"""
    add_omml_equation(doc, OMML_CR_ER)

    add_legend(doc, [
        ("r₀",   "rayon initial de l'élément sensible (mm)"),
        ("r(t)", "rayon à l'instant t (mm), calculé depuis R(t)"),
        ("t",    "durée écoulée (heures)"),
        ("8760", "facteur de conversion heures → années"),
    ])

    add_heading(doc, "2.4.2  Sondes ER commerciales et chaîne de mesure industrielle", 3)

    add_body(doc,
        "Les principaux fabricants de sondes ER pour l'industrie pétrolière sont Emerson "
        "(gamme Roxar), Cormon Ltd, et Permasense (sondes permanentes sans câble). Ces "
        "équipements se déclinent en plusieurs géométries d'éléments sensibles, choisies "
        "selon la sensibilité requise et la durée de vie souhaitée :"
    )

    sonde_data = [
        ("Élément fil (wire)",    "Fil circulaire exposé — haute sensibilité, courte durée de vie",    "Mesures courte durée, laboratoire"),
        ("Élément tube (tube)",   "Tube cylindrique — bonne sensibilité, durée de vie moyenne",        "Production standard"),
        ("Élément flush",         "Plaquette affleurante — faible sensibilité, très longue durée",     "Surveillance long terme"),
        ("Élément ruban (strip)", "Ruban plat — sensibilité intermédiaire",                            "Applications offshore"),
    ]
    table_sonde = doc.add_table(rows=1, cols=3)
    table_sonde.style = 'Table Grid'
    for i, txt in enumerate(["Type", "Caractéristiques", "Application"]):
        table_sonde.rows[0].cells[i].text = txt
        table_sonde.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_sonde.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_sonde.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    for row_d in sonde_data:
        row = table_sonde.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_body(doc,
        "La chaîne de mesure industrielle complète comprend : la sonde ER (élément sensible) "
        "→ le transmetteur (convertisseur R → signal 4-20 mA ou numérique) → le réseau de "
        "terrain (HART, Modbus, WirelessHART) → le système d'historisation (PI Server "
        "OSIsoft/AVEVA) → le SCADA ou DCS. Le prix d'une sonde ER industrielle complète "
        "varie entre 2 000 et 15 000 USD selon le type et la connectivité."
    )

    add_heading(doc, "2.5  Méthodes non-destructives (CND)", 2)

    add_body(doc,
        "Les méthodes de contrôle non-destructif (CND) mesurent l'épaisseur résiduelle de "
        "paroi sans consommer l'élément sensible. Elles sont complémentaires des sondes ER "
        "car elles donnent une mesure absolue de l'état de la structure, là où la sonde ER "
        "donne une vitesse de corrosion instantanée."
    )

    add_body(doc,
        "La mesure d'épaisseur par ultrasons (UT classique, norme ASTM E797) utilise la "
        "propagation d'une onde ultrasonore dans le métal et la mesure du temps de retour "
        "de l'écho de fond. La précision typique est de ±0,1 mm, suffisante pour détecter "
        "des pertes d'épaisseur significatives. La technique PAUT (Phased Array UT) permet "
        "une imagerie 2D ou 3D de la corrosion en déplaçant électroniquement le faisceau "
        "ultrasonore, offrant une cartographie complète de l'état de la paroi. Le Guided "
        "Wave Testing (GWT) propage des ondes guidées sur de longues distances (jusqu'à "
        "100 m) depuis un seul point d'accès, permettant l'inspection de tronçons entiers "
        "de pipeline sans excavation."
    )

    add_heading(doc, "2.6  Synthèse comparative et justification du choix ER", 2)

    table_comp = doc.add_table(rows=1, cols=5)
    table_comp.style = 'Table Grid'
    for i, txt in enumerate(["Méthode", "Continu", "Low-cost", "Précision", "Adapté labo DIY"]):
        table_comp.rows[0].cells[i].text = txt
        table_comp.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_comp.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_comp.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    comp_data = [
        ("Gravimétrie",   "Non",  "Oui",  "Bonne",      "Oui (référence)"),
        ("LPR",           "Oui",  "Non",  "Bonne",      "Difficile (potentiostat)"),
        ("EIS",           "Non",  "Non",  "Très bonne", "Non (équipement spécialisé)"),
        ("Sonde ER",      "Oui",  "Oui*", "Moyenne",    "Oui (fil + Wheatstone)"),
        ("UT classique",  "Non",  "Non",  "Bonne",      "Non (transducteur piézo)"),
        ("PAUT",          "Non",  "Non",  "Très bonne", "Non"),
        ("GWT",           "Non",  "Non",  "Moyenne",    "Non"),
    ]
    for row_d in comp_data:
        row = table_comp.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_body(doc,
        "* Low-cost uniquement pour la version DIY développée dans ce travail. "
        "Les sondes ER commerciales coûtent 2 000–15 000 USD."
    )

    add_body(doc,
        "La sonde ER est la seule méthode combinant surveillance continue, mesure directe "
        "du taux de corrosion et transposabilité en laboratoire à faible coût. La "
        "gravimétrie sera utilisée en parallèle comme méthode de référence pour valider "
        "les mesures de la sonde DIY : les coupons exposés dans le même bac permettront "
        "de vérifier la cohérence entre la perte de masse mesurée et le CR calculé "
        "depuis la résistance électrique."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PARTIE 3
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PARTIE 3 — DE LA SONDE ER INDUSTRIELLE AU PROTOTYPE DE LABORATOIRE", 1)

    add_heading(doc, "3.1  Analogie composant par composant", 2)

    add_body(doc,
        "L'objectif de cette partie est de démontrer que le prototype développé dans ce "
        "travail n'est pas un système ad hoc sans référence industrielle, mais une "
        "transposition rigoureuse et délibérée d'une chaîne de mesure ER industrielle "
        "normalisée (ASTM G96) vers un équivalent de laboratoire low-cost. Chaque "
        "composant industriel a son équivalent fonctionnel dans le prototype :"
    )

    table_ana = doc.add_table(rows=1, cols=4)
    table_ana.style = 'Table Grid'
    for i, txt in enumerate(["Fonction", "Composant industriel", "Équivalent prototype", "Principe identique ?"]):
        table_ana.rows[0].cells[i].text = txt
        table_ana.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_ana.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_ana.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    ana_data = [
        ("Élément sensible corrodable", "Fil/tube acier certifié (Cormon, Emerson)", "Fil de fer 0,2 mm", "Oui — même loi R=ρL/A"),
        ("Mesure différentielle", "Pont de Wheatstone intégré dans transmetteur", "Pont de Wheatstone discret (résistances 10Ω)", "Oui — même circuit"),
        ("Amplification signal", "Amplificateur de précision (INA128, AD8221)", "ADS1115 (PGA intégré, 16 bits)", "Oui — même principe PGA"),
        ("Traitement numérique", "Microprocesseur embarqué dans transmetteur", "ESP8266 (microcontrôleur WiFi)", "Oui — même fonction"),
        ("Communication", "HART 4-20 mA, WirelessHART, Modbus", "HTTP/REST via WiFi 802.11", "Équivalent — protocoles différents"),
        ("Historisation", "PI Server OSIsoft (historien industriel)", "Supabase (PostgreSQL cloud)", "Équivalent — stockage temporel"),
        ("Visualisation", "SCADA/DCS (Wonderware, Honeywell)", "Dashboard Streamlit", "Équivalent — affichage temps réel"),
        ("Prédiction / Décision", "Absent dans systèmes classiques", "XGBoost (contribution originale)", "Nouveau — apport de ce travail"),
    ]
    for row_d in ana_data:
        row = table_ana.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "3.2  Conception du pont de Wheatstone", 2)

    add_body(doc,
        "Le pont de Wheatstone est un circuit électrique permettant de mesurer avec précision "
        "de très faibles variations de résistance en mode différentiel, annulant ainsi le "
        "bruit commun (variations de température, bruit d'alimentation). Il est constitué "
        "de quatre résistances disposées en carré, avec une source de tension Vs et une "
        "mesure de tension différentielle Vout entre les deux branches :"
    )

    add_body(doc,
        "Dans la configuration utilisée, R1 = R2 = R3 = R₀ (résistances fixes de précision "
        "10Ω, tolérance 1%) et R4 = R_fil (sonde ER, résistance variable). À l'équilibre "
        "initial (t=0, fil non corrodé), R_fil = R₀ et Vout = 0. Lorsque le fil corrode, "
        "R_fil > R₀ et la tension de déséquilibre Vout est donnée par :"
    )

    OMML_WHEAT = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>V</m:t></m:r>
      <m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>out</m:t></m:r></m:sub></m:sSub>
      <m:r><m:t>&#160;=&#160;V</m:t></m:r>
      <m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>s</m:t></m:r></m:sub></m:sSub>
      <m:r><m:t>&#160;&#xD7;&#160;</m:t></m:r>
      <m:d>
        <m:dPr><m:begChr m:val="("/>  <m:endChr m:val=")"/></m:dPr>
        <m:e>
          <m:f>
            <m:num><m:r><m:t>R</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>fil</m:t></m:r></m:sub></m:sSub></m:num>
            <m:den><m:r><m:t>R&#x2080;&#160;+&#160;R</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>fil</m:t></m:r></m:sub></m:sSub></m:den>
          </m:f>
          <m:r><m:t>&#160;&#x2212;&#160;</m:t></m:r>
          <m:f>
            <m:num><m:r><m:t>R&#x2080;</m:t></m:r></m:num>
            <m:den><m:r><m:t>2&#160;&#xB7;&#160;R&#x2080;</m:t></m:r></m:den>
          </m:f>
        </m:e>
      </m:d>
    </m:oMath>"""
    add_omml_equation(doc, OMML_WHEAT)

    add_body(doc,
        "L'ADS1115 mesure Vout en mode différentiel (AIN0 - AIN1) avec un gain PGA de ×16, "
        "ce qui donne une résolution de 0,0078 mV par bit sur une plage de ±256 mV. "
        "La résolution en résistance correspondante est de l'ordre de 0,001 Ω, suffisante "
        "pour détecter des variations de taux de corrosion supérieures à 0,01 mm/an."
    )

    add_heading(doc, "3.3  Environnement corrosif contrôlé et justification de l'analogie", 2)

    add_body(doc,
        "L'environnement corrosif du prototype (acide acétique CH₃COOH 6–14% + NaCl "
        "20–70 g/L) est une approximation de laboratoire des conditions rencontrées dans "
        "les pipelines Oil & Gas. La justification de cette analogie repose sur la "
        "communauté du mécanisme électrochimique fondamental :"
    )

    ana2_data = [
        ("Milieu acide pétrolier", "CO₂ + H₂O → H₂CO₃ → H⁺ + HCO₃⁻\nRéaction cathodique : 2H⁺ + 2e⁻ → H₂"),
        ("Milieu acide prototype", "CH₃COOH → CH₃COO⁻ + H⁺\nRéaction cathodique : 2H⁺ + 2e⁻ → H₂"),
    ]

    for label, desc in ana2_data:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{label} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11)
        set_paragraph_spacing(p, before=0, after=4)

    add_body(doc,
        "Dans les deux cas, la réaction cathodique est la réduction des protons H⁺. "
        "L'acide acétique est un acide faible (pKa = 4,76) qui fournit des protons de "
        "manière contrôlée, contrairement aux acides forts (HCl, H₂SO₄) qui donneraient "
        "une corrosion trop rapide et difficile à contrôler en laboratoire. Les ions "
        "chlorures (Cl⁻) du NaCl jouent le même rôle agressif qu'en milieu pétrolier : "
        "destruction du film passif et accélération de la corrosion par piqûres."
    )

    add_body(doc,
        "Les limites de cette analogie sont clairement identifiées : l'acide acétique "
        "ne génère pas de CO₂ dissous et ne reproduit pas la corrosion à haute pression "
        "des pipelines profonds. Ces conditions seraient reproductibles uniquement en "
        "autoclave pressurisé, équipement hors budget. L'analogie est donc qualitative "
        "et suffisante pour démontrer les principes du système de surveillance et de "
        "prédiction, objectif central de ce mémoire."
    )

    add_heading(doc, "3.4  Protocole expérimental — 4 phases", 2)

    add_body(doc,
        "Le protocole expérimental est conçu selon un plan factoriel partiel à 4 phases "
        "successives de 7 jours chacune, permettant d'isoler l'effet de chaque variable "
        "de protection sur le taux de corrosion :"
    )

    phase_data = [
        ("Phase 1 — Baseline (J0–J7)",
         "NaCl 20 g/L + vinaigre 6%, T=30°C, sans inhibiteur, sans protection.\n"
         "Objectif : établir le taux de corrosion de référence sans aucune protection."),
        ("Phase 2 — Inhibiteur (J7–J14)",
         "NaCl 50 g/L + vinaigre 14%, T=30°C, inhibiteur huile végétale 1 mL/L.\n"
         "Objectif : quantifier l'efficacité inhibitrice η et entraîner le modèle "
         "à reconnaître la signature de l'inhibiteur."),
        ("Phase 3 — Protection cathodique (J14–J21)",
         "NaCl 50 g/L, T=50°C, pompe à air (aération), anode sacrificielle zinc.\n"
         "Objectif : reproduire les conditions industrielles agressives et mesurer "
         "l'effet de la protection cathodique."),
        ("Phase 4 — Combinaison optimale (J21–J28)",
         "NaCl 50 g/L, T=50°C, inhibiteur 2 mL/L + protection cathodique.\n"
         "Objectif : identifier les conditions minimisant le CR — données pour "
         "l'optimisation de la recommandation XGBoost."),
    ]

    for titre, desc in phase_data:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{titre} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=6)

    add_body(doc,
        "Les mesures sont effectuées toutes les 10 minutes (144 mesures/jour, "
        "soit ~4 032 points sur 28 jours). La fréquence de 10 minutes est un compromis "
        "entre la résolution temporelle et la consommation énergétique de l'ESP8266 "
        "(mode deep sleep entre les mesures)."
    )

    add_heading(doc, "3.5  Limites du prototype par rapport au système industriel", 2)

    limites = [
        ("Certification",       "Le prototype n'est pas certifié ATEX ni IECEx — inutilisable en zone explosive réelle."),
        ("Pression",            "Mesure à pression atmosphérique uniquement. Les pipelines opèrent à 50–200 bar."),
        ("Environnement",       "Acide acétique ≠ CO₂/H₂S dissous — analogie qualitative uniquement."),
        ("Durée de vie",        "Fil de fer 0,2 mm : durée de vie estimée 7–14 jours selon CR. Sonde industrielle : 1–5 ans."),
        ("Résolution ADC",      "ADS1115 16 bits vs convertisseurs industriels 24 bits — résolution 10× inférieure."),
        ("Redondance",          "Pas de système de sauvegarde ni de watchdog industriel."),
    ]

    for limite, desc in limites:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{limite} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=4)

    add_note(doc,
        "Ces limites sont assumées et documentées. L'objectif du prototype n'est pas de "
        "remplacer un système industriel certifié, mais de démontrer la faisabilité de "
        "l'approche ML+IoT pour la maintenance prédictive de la corrosion, avec des "
        "moyens accessibles dans un contexte académique."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PARTIE 4
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PARTIE 4 — MAINTENANCE PRÉDICTIVE ET MACHINE LEARNING", 1)

    add_heading(doc, "4.1  Évolution des stratégies de maintenance industrielle", 2)

    add_body(doc,
        "La maintenance industrielle a connu une évolution continue depuis les débuts de "
        "l'industrialisation, passant d'une approche purement réactive à une approche "
        "de plus en plus anticipative. Quatre générations de stratégies se sont "
        "successivement développées :"
    )

    maint_data = [
        ("Maintenance corrective (1ère génération)",
         "Intervention après défaillance. Aucun investissement préventif, mais coûts de "
         "défaillance élevés et imprévisibles. Toujours pertinente pour les équipements "
         "non critiques (norme EN 13306)."),
        ("Maintenance préventive systématique (2ème génération)",
         "Remplacement des composants à intervalles fixes, indépendamment de leur état réel. "
         "Sûre mais économiquement inefficace : on remplace des pièces encore en bon état."),
        ("Maintenance conditionnelle (3ème génération)",
         "Intervention déclenchée par le dépassement d'un seuil mesuré (vibration, température, "
         "épaisseur). Représente le standard industriel actuel pour les équipements critiques."),
        ("Maintenance prédictive (4ème génération)",
         "Prédiction de la défaillance future par modèles mathématiques ou ML, permettant "
         "d'intervenir au moment optimal (ni trop tôt, ni trop tard). Objet de ce travail."),
    ]

    for titre, desc in maint_data:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{titre} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=6)

    add_heading(doc, "4.2  Inspection Basée sur le Risque (RBI) — API 580/581", 2)

    add_body(doc,
        "L'API 580/581 (Risk-Based Inspection) est la norme de référence internationale "
        "pour l'optimisation des plans d'inspection des équipements sous pression dans "
        "l'industrie pétrolière et chimique. Elle définit le risque comme le produit de "
        "la probabilité de défaillance (PoF) par la conséquence de la défaillance (CoF) :"
    )

    OMML_RISK = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>Risque&#160;=&#160;PoF&#160;&#xD7;&#160;CoF</m:t></m:r>
    </m:oMath>"""
    add_omml_equation(doc, OMML_RISK)

    add_body(doc,
        "La durée de vie résiduelle (RUL — Remaining Useful Life) d'un équipement soumis "
        "à la corrosion est calculée selon l'ASME B31.3 à partir de l'épaisseur mesurée "
        "et du taux de corrosion :"
    )

    OMML_RUL = f"""<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>RUL&#160;(ans)&#160;=&#160;</m:t></m:r>
      <m:f>
        <m:num><m:r><m:t>e</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>mesur&#233;</m:t></m:r></m:sub></m:sSub><m:r><m:t>&#160;&#x2212;&#160;e</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>min</m:t></m:r></m:sub></m:sSub></m:num>
        <m:den><m:r><m:t>CR</m:t></m:r></m:den>
      </m:f>
    </m:oMath>"""
    add_omml_equation(doc, OMML_RUL)

    add_legend(doc, [
        ("e_mesuré", "épaisseur de paroi mesurée (mm)"),
        ("e_min",    "épaisseur minimale admissible selon ASME B31.3 (mm)"),
        ("CR",       "taux de corrosion prédit par XGBoost (mm/an)"),
    ])

    add_body(doc,
        "La contribution principale de ce travail à la méthodologie RBI est de remplacer "
        "le CR mesuré historiquement (valeur passée) par un CR prédit par XGBoost (valeur "
        "future), ce qui permet de calculer un RUL prospectif plutôt que rétrospectif, "
        "offrant ainsi une anticipation réelle des interventions de maintenance."
    )

    add_heading(doc, "4.3  État de l'art — Machine Learning appliqué à la corrosion (2015–2024)", 2)

    add_body(doc,
        "L'application du machine learning à la prédiction de la corrosion est un domaine "
        "de recherche en croissance rapide. Une revue de la littérature des dix dernières "
        "années permet d'identifier les approches dominantes et leurs performances comparatives."
    )

    add_body(doc,
        "Les premières applications ML à la corrosion ont utilisé les réseaux de neurones "
        "artificiels (ANN) pour modéliser la relation non-linéaire entre les paramètres "
        "physico-chimiques et le taux de corrosion (Caleyo et al., 2009 ; Ossai et al., 2015). "
        "Ces modèles ont démontré une supériorité par rapport aux modèles physiques purs "
        "(De Waard-Milliams) en présence de données hétérogènes, mais souffrent d'un "
        "manque d'interprétabilité — problème critique pour les ingénieurs de corrosion "
        "qui doivent justifier leurs décisions d'inspection."
    )

    add_body(doc,
        "Les forêts aléatoires (Random Forest) et les machines à vecteurs de support "
        "(SVM) ont ensuite été appliquées avec succès à la classification du risque de "
        "corrosion et à la prédiction des défaillances de pipelines (Xie et Tian, 2018 ; "
        "Almutairi et al., 2020). Ces méthodes offrent une meilleure robustesse face au "
        "bruit de mesure, mais leurs performances sur données tabulaires sont généralement "
        "inférieures au gradient boosting."
    )

    add_body(doc,
        "Les modèles LSTM (Long Short-Term Memory) ont été proposés pour exploiter la "
        "structure temporelle des séries de mesures de corrosion (Pidaparti et al., 2021). "
        "Bien que théoriquement adaptés aux données séquentielles, ils nécessitent un "
        "volume de données important (>10 000 points) et présentent une sensibilité élevée "
        "aux hyperparamètres, ce qui limite leur applicabilité dans des contextes à données "
        "restreintes comme ce prototype."
    )

    add_body(doc,
        "Le gradient boosting, et particulièrement XGBoost (Chen et Guestrin, 2016), "
        "s'est imposé comme la méthode de référence pour la prédiction sur données tabulaires "
        "industrielles dans de nombreux benchmarks (Kaggle, études comparatives). Son "
        "application à la prédiction de corrosion a été validée par plusieurs études récentes "
        "(Li et al., 2022 ; Soomro et al., 2023) qui rapportent des performances supérieures "
        "aux ANN et RF sur des jeux de données de taille limitée (<5 000 points)."
    )

    add_heading(doc, "4.4  XGBoost — Fondements et mécanisme de décision", 2)

    # 4.4.1 Arbre de décision
    add_heading(doc, "4.4.1  L'arbre de décision : brique élémentaire", 3)

    add_body(doc,
        "Tout algorithme de gradient boosting repose sur un modèle de base : l'arbre de "
        "décision. Un arbre de décision est une structure hiérarchique de règles binaires "
        "appliquées successivement aux variables d'entrée (features) pour aboutir à une "
        "prédiction. Chaque nœud interne teste une condition sur une feature "
        "(par exemple : delta_R_1h > 0,004 ?) et oriente l'observation vers la branche "
        "gauche (condition vraie) ou droite (condition fausse). Les feuilles terminales "
        "contiennent la valeur prédite, calculée comme la moyenne des observations "
        "d'entraînement qui ont atteint cette feuille."
    )

    add_body(doc,
        "L'algorithme d'apprentissage d'un arbre cherche, à chaque nœud, le seuil de "
        "séparation qui minimise la variance des résidus dans chaque sous-groupe résultant "
        "(critère de réduction de variance pour la régression). Ce processus est répété "
        "récursivement jusqu'à atteindre une profondeur maximale (max_depth) ou un nombre "
        "minimal d'observations par feuille (min_child_weight)."
    )

    add_body(doc,
        "La limitation fondamentale d'un arbre unique réside dans son compromis biais-variance : "
        "un arbre peu profond présente un biais élevé (sous-apprentissage — il ne capture pas "
        "les nuances des données) ; un arbre très profond présente une variance élevée "
        "(sur-apprentissage — il mémorise le bruit de mesure). Le boosting résout ce dilemme "
        "en combinant un grand nombre d'arbres peu profonds."
    )

    # 4.4.2 Boosting
    add_heading(doc, "4.4.2  Le boosting : correction itérative des erreurs", 3)

    add_body(doc,
        "Le boosting est une méthode d'apprentissage ensembliste qui construit séquentiellement "
        "une série d'apprenants faibles (arbres peu profonds), chaque apprenant ciblant "
        "spécifiquement les erreurs commises par l'ensemble des apprenants précédents. "
        "Formellement, la prédiction finale est une somme pondérée de K arbres :"
    )

    OMML_BOOST = """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>&#x177;&#160;=&#160;</m:t></m:r>
      <m:nary>
        <m:naryPr>
          <m:chr m:val="&#x2211;"/>
          <m:subHide m:val="0"/>
          <m:supHide m:val="0"/>
        </m:naryPr>
        <m:sub><m:r><m:t>k=1</m:t></m:r></m:sub>
        <m:sup><m:r><m:t>K</m:t></m:r></m:sup>
        <m:e><m:r><m:t>&#x3B7;&#160;&#xB7;&#160;f&#x1D4F;(x)</m:t></m:r></m:e>
      </m:nary>
    </m:oMath>"""
    add_omml_equation(doc, OMML_BOOST)

    add_legend(doc, [
        ("ŷ",    "prédiction finale (taux de corrosion en mm/an)"),
        ("fk(x)", "prédiction du k-ième arbre pour l'observation x"),
        ("η",    "learning rate — contrôle la contribution de chaque arbre (typiquement 0,05–0,3)"),
        ("K",    "nombre total d'arbres (n_estimators, typiquement 100–500)"),
    ])

    add_body(doc,
        "Le processus d'entraînement procède comme suit. À l'initialisation, la prédiction "
        "courante ŷ⁽⁰⁾ est fixée à la moyenne de la variable cible sur l'ensemble "
        "d'entraînement. À chaque itération k, le résidu rᵢ = yᵢ − ŷᵢ⁽ᵏ⁻¹⁾ est calculé "
        "pour chaque observation. Un nouvel arbre fk est entraîné à prédire ces résidus. "
        "La prédiction est ensuite mise à jour : ŷᵢ⁽ᵏ⁾ = ŷᵢ⁽ᵏ⁻¹⁾ + η · fk(xᵢ). "
        "L'arbre k ne prédit donc pas le taux de corrosion directement — il prédit "
        "l'écart entre la réalité et ce que les arbres précédents ont déjà estimé."
    )

    add_body(doc,
        "À titre d'illustration sur les données de corrosion : supposons que la prédiction "
        "initiale soit ŷ⁽⁰⁾ = 0,16 mm/an (moyenne de la phase de référence). Pour une "
        "observation réelle CR = 0,31 mm/an (phase corrosion intense), le résidu est "
        "r = +0,15. L'Arbre 1 apprend que lorsque delta_R_1h est élevé et la température "
        "dépasse 48°C, le résidu est fortement positif. Il contribue +0,12. L'Arbre 2 "
        "capture le résidu résiduel +0,03, et ainsi de suite jusqu'à convergence."
    )

    # 4.4.3 Gradient boosting
    add_heading(doc, "4.4.3  Du boosting au gradient boosting", 3)

    add_body(doc,
        "Friedman (2001) a généralisé le boosting en formulant chaque itération comme une "
        "étape de descente de gradient dans l'espace fonctionnel. L'idée centrale est que "
        "le résidu rᵢ = yᵢ − ŷᵢ est mathématiquement équivalent au gradient négatif de "
        "la fonction de perte MSE par rapport à la prédiction courante :"
    )

    OMML_GRAD = """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>r&#x1D35;&#160;=&#160;&#x2212;</m:t></m:r>
      <m:f>
        <m:num>
          <m:r><m:t>&#x2202;L(y&#x1D35;,&#160;&#x177;&#x1D35;)</m:t></m:r>
        </m:num>
        <m:den>
          <m:r><m:t>&#x2202;&#x177;&#x1D35;</m:t></m:r>
        </m:den>
      </m:f>
      <m:r><m:t>&#160;=&#160;y&#x1D35;&#160;&#x2212;&#160;&#x177;&#x1D35;&#160;&#160;&#160;&#160;(pour L&#160;=&#160;MSE)</m:t></m:r>
    </m:oMath>"""
    add_omml_equation(doc, OMML_GRAD)

    add_body(doc,
        "Cette reformulation est décisive car elle permet de substituer la MSE par toute "
        "fonction de perte différentiable (MAE, Huber, quantile), chaque choix modifiant "
        "la sensibilité du modèle aux valeurs extrêmes. En pratique pour la prédiction du "
        "taux de corrosion, la MSE est préférée car elle pénalise davantage les grandes "
        "erreurs — une sous-estimation sévère du taux de corrosion étant plus dangereuse "
        "qu'une légère surestimation."
    )

    # 4.4.4 Spécificités XGBoost
    add_heading(doc, "4.4.4  Apports spécifiques de XGBoost", 3)

    add_body(doc,
        "XGBoost (Chen et Guestrin, 2016) améliore le gradient boosting classique sur "
        "trois points fondamentaux qui le rendent particulièrement robuste dans les "
        "contextes industriels à données limitées et bruitées."
    )

    add_body(doc,
        "Premièrement, la fonction objectif intègre un terme de régularisation explicite "
        "qui pénalise la complexité de chaque arbre :"
    )

    OMML_XGB = """<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
      <m:r><m:t>Obj&#160;=&#160;</m:t></m:r>
      <m:nary>
        <m:naryPr>
          <m:chr m:val="&#x2211;"/>
          <m:subHide m:val="0"/>
          <m:supHide m:val="0"/>
        </m:naryPr>
        <m:sub><m:r><m:t>i</m:t></m:r></m:sub>
        <m:sup><m:r><m:t>n</m:t></m:r></m:sup>
        <m:e><m:r><m:t>L(y&#x1D35;,&#160;&#x177;&#x1D35;)&#160;+&#160;</m:t></m:r>
          <m:nary>
            <m:naryPr><m:chr m:val="&#x2211;"/><m:subHide m:val="0"/><m:supHide m:val="0"/></m:naryPr>
            <m:sub><m:r><m:t>k</m:t></m:r></m:sub>
            <m:sup><m:r><m:t>K</m:t></m:r></m:sup>
            <m:e><m:r><m:t>&#x3A9;(f&#x1D4F;)</m:t></m:r></m:e>
          </m:nary>
        </m:e>
      </m:nary>
    </m:oMath>"""
    add_omml_equation(doc, OMML_XGB)

    add_legend(doc, [
        ("L(yi, ŷi)", "fonction de perte — MSE : (yi − ŷi)²"),
        ("Ω(fk)",     "regularisation : γT + ½λ||w||²  (T = nb feuilles, w = poids feuilles)"),
        ("γ",         "pénalise le nombre de feuilles — contrôle la complexité structurelle"),
        ("λ",         "régularisation L2 sur les poids — réduit la sensibilité au bruit"),
    ])

    add_body(doc,
        "Le paramètre γ pénalise le nombre de feuilles de chaque arbre : un arbre avec "
        "beaucoup de feuilles n'est retenu que si le gain en précision justifie sa complexité. "
        "Le paramètre λ applique une régularisation L2 sur les poids des feuilles, évitant "
        "que le modèle s'ajuste aux fluctuations parasites du signal de résistance "
        "(bruit électronique de l'ADS1115, variations de connectivité)."
    )

    add_body(doc,
        "Deuxièmement, XGBoost implémente un élagage post-construction (pruning) basé sur "
        "un gain seuil. Contrairement au gradient boosting classique qui arrête la croissance "
        "de l'arbre dès qu'une division est jugée mauvaise, XGBoost construit l'arbre jusqu'à "
        "la profondeur maximale puis élague les branches dont le gain net est inférieur à γ. "
        "Cette stratégie globale évite de rejeter des divisions localement mauvaises mais "
        "globalement bénéfiques."
    )

    add_body(doc,
        "Troisièmement, XGBoost gère nativement les valeurs manquantes en apprenant "
        "automatiquement, à chaque nœud, dans quelle branche orienter les observations "
        "présentant une valeur NaN. Dans le contexte de ce prototype, des interruptions de "
        "connectivité WiFi de l'ESP8266 peuvent créer des lacunes dans les séries temporelles. "
        "XGBoost apprend que ces lacunes correspondent généralement à un contexte particulier "
        "(perte de signal en phase de forte corrosion, par exemple) et traite les NaN de "
        "manière cohérente plutôt que de rejeter ces observations."
    )

    add_body(doc,
        "Ces trois mécanismes combinés confèrent à XGBoost une robustesse particulière "
        "dans les contextes à données limitées (<5 000 points) et bruitées, caractéristiques "
        "typiques des expérimentations en laboratoire. Chen et Guestrin (2016) ont validé "
        "ces propriétés sur 29 benchmarks de compétitions de machine learning, établissant "
        "XGBoost comme la méthode de référence pour les données tabulaires industrielles. "
        "Cette supériorité a été confirmée dans le domaine de la corrosion par Li et al. "
        "(2022) et Soomro et al. (2023), qui rapportent des R² supérieurs à 0,92 sur des "
        "jeux de données de moins de 3 000 points, nettement supérieurs aux ANN (R² ~ 0,85) "
        "et aux forêts aléatoires (R² ~ 0,88) dans les mêmes conditions."
    )

    add_note(doc,
        "Dans ce travail, XGBoost est entraîné en mode régression pour prédire le taux "
        "de corrosion CR (mm/an) à horizon 48h. La recommandation de dose d'inhibiteur "
        "est ensuite déterminée par un moteur de règles post-modèle (CR < 0,05 mm/an → "
        "aucune action ; 0,05–0,15 → 0,5 mL/L ; 0,15–0,30 → 1 mL/L ; > 0,30 → 2 mL/L + alerte), "
        "découplant la prédiction probabiliste de la décision opérationnelle."
    )

    add_heading(doc, "4.5  Feature engineering pour séries temporelles de corrosion", 2)

    add_body(doc,
        "La mesure brute de résistance toutes les 10 minutes présente une forte "
        "autocorrélation temporelle : deux mesures consécutives sont quasi identiques "
        "car la corrosion est un processus lent. Fournir ces mesures brutes à XGBoost "
        "reviendrait à lui présenter ~4 000 observations quasi identiques, sans information "
        "sur la dynamique du phénomène. Le feature engineering transforme ces données "
        "brutes en features porteuses d'information temporelle :"
    )

    fe_data = [
        ("delta_R_1h",       "R(t) - R(t-6 mesures)",          "Vitesse de corrosion sur la dernière heure"),
        ("delta_R_6h",       "R(t) - R(t-36 mesures)",         "Tendance sur 6 heures"),
        ("delta_R_24h",      "R(t) - R(t-144 mesures)",        "Tendance journalière"),
        ("vitesse_CR_1h",    "delta_R_1h / (π·r·L)",           "Taux de corrosion instantané (mm/an)"),
        ("tendance_24h",     "Pente linéaire sur 144 mesures",  "Accélération ou décélération"),
        ("temp_moy_6h",      "Moyenne glissante temp sur 6h",   "Température stable filtrée"),
        ("dose_inhibiteur",  "Dose injectée (mL/L)",            "Variable de contrôle — protection"),
        ("protection_cp",    "Booléen 0/1",                     "Présence anode sacrificielle"),
    ]

    table_fe = doc.add_table(rows=1, cols=3)
    table_fe.style = 'Table Grid'
    for i, txt in enumerate(["Feature", "Calcul", "Information portée"]):
        table_fe.rows[0].cells[i].text = txt
        table_fe.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_fe.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_fe.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    for row_d in fe_data:
        row = table_fe.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "4.6  Validation temporelle — Walk-forward Cross-Validation", 2)

    add_body(doc,
        "La validation croisée standard (k-fold aléatoire) est inadaptée aux séries "
        "temporelles car elle introduit une fuite temporelle (data leakage) : des données "
        "futures peuvent se retrouver dans l'ensemble d'entraînement, ce qui biaise "
        "artificiellement les métriques à la hausse. Pour les données de corrosion, "
        "la méthode appropriée est la walk-forward cross-validation (ou TimeSeriesSplit) :"
    )

    add_body(doc,
        "Dans cette approche, les données sont divisées en N folds chronologiques. "
        "À chaque itération k, le modèle est entraîné sur les folds 1 à k et évalué "
        "sur le fold k+1 (données strictement futures). La performance finale est la "
        "moyenne des métriques sur tous les folds de validation. Cette méthode garantit "
        "que le modèle n'a jamais accès à des données futures lors de l'entraînement, "
        "ce qui reflète fidèlement les conditions d'utilisation réelle."
    )

    add_heading(doc, "4.7  Boucle fermée : mesure → prédiction → recommandation", 2)

    add_body(doc,
        "Le système développé dans ce travail implémente une boucle de contrôle prédictive "
        "fermée, reliant directement la mesure en temps réel à la recommandation d'action :"
    )

    boucle = [
        ("Acquisition",    "ESP8266 + ADS1115 mesurent R, T, pH toutes les 10 minutes → Supabase"),
        ("Feature calc.",  "Pipeline Python calcule delta_R_1h, delta_R_24h, tendance_24h..."),
        ("Prédiction",     "XGBoost prédit CR dans les 48 prochaines heures"),
        ("Décision",       "Moteur de règles : CR < 0,05 → OK ; 0,05–0,15 → 1 mL/L ; > 0,15 → alerte + 2 mL/L"),
        ("Action",         "Recommandation affichée sur dashboard Streamlit + alerte optionnelle"),
        ("Feedback",       "La dose injectée devient une feature pour le prochain cycle d'entraînement"),
    ]

    for i, (step, desc) in enumerate(boucle, 1):
        p = doc.add_paragraph(style='List Number')
        r1 = p.add_run(f"{step} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=4)

    add_note(doc,
        "Cette boucle fermée constitue l'apport original de ce travail par rapport aux "
        "systèmes ER classiques : ceux-ci mesurent et enregistrent, mais ne prédisent "
        "pas et ne recommandent pas. L'intégration du ML ferme la boucle entre la "
        "surveillance et l'action de maintenance."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PARTIE 5
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "PARTIE 5 — POSITIONNEMENT ET CONTRIBUTION ORIGINALE", 1)

    add_heading(doc, "5.1  Gaps identifiés dans la littérature", 2)

    add_body(doc,
        "L'analyse de la littérature présentée dans les parties précédentes permet "
        "d'identifier trois gaps principaux que ce travail vise à adresser :"
    )

    gaps = [
        ("Gap 1 — Accessibilité des systèmes de surveillance",
         "Les systèmes de surveillance ER industriels (Cormon, Emerson, Permasense) "
         "coûtent entre 2 000 et 15 000 USD par point de mesure, rendant leur déploiement "
         "inaccessible aux PME industrielles et aux économies émergentes. Aucune solution "
         "low-cost (<100 USD) combinant mesure ER continue et transmission IoT n'a été "
         "documentée dans la littérature académique."),
        ("Gap 2 — Prédiction vs. Surveillance",
         "La majorité des systèmes ER industriels sont des outils de surveillance : "
         "ils mesurent et enregistrent le taux de corrosion passé mais ne prédisent "
         "pas son évolution future. Les travaux ML sur la corrosion utilisent presque "
         "exclusivement des données publiques (PHMSA, SPE papers) générées dans des "
         "contextes industriels sans accès direct au système de mesure. La boucle "
         "mesure → prédiction → recommandation en temps réel n'est pas documentée "
         "à l'échelle d'un prototype fonctionnel complet."),
        ("Gap 3 — Contexte africain",
         "La littérature sur la maintenance prédictive de la corrosion dans le contexte "
         "des industries extractives africaines (Oil & Gas, mines) est quasi inexistante. "
         "Les solutions développées pour les environnements nord-américains ou européens "
         "ne prennent pas en compte les contraintes locales : connectivité intermittente, "
         "budget limité, disponibilité des pièces, expertise technique restreinte."),
    ]

    for titre, desc in gaps:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{titre} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=6)

    add_heading(doc, "5.2  Contribution originale de ce travail", 2)

    add_body(doc,
        "Face aux gaps identifiés, ce travail apporte les contributions suivantes :"
    )

    contribs = [
        ("Contribution 1 — Prototype IoT low-cost",
         "Conception et implémentation d'une chaîne de mesure ER complète (sonde DIY + "
         "ADS1115 + ESP8266 + Supabase) pour un coût total de hardware inférieur à "
         "30 000 FCFA (~50 USD), soit un facteur 40 à 300 moins cher qu'un système "
         "industriel certifié, tout en conservant le même principe de mesure normalisé (ASTM G96)."),
        ("Contribution 2 — Intégration ML prédictif",
         "Premier prototype documenté intégrant une chaîne complète mesure ER → feature "
         "engineering temporel → XGBoost → recommandation de dose d'inhibiteur en temps "
         "réel, fermant la boucle entre la surveillance et la décision de maintenance."),
        ("Contribution 3 — Validation expérimentale",
         "Production de données de corrosion réelles en laboratoire contrôlé "
         "(4 phases × 7 jours, 144 mesures/jour) servant à entraîner et valider le modèle "
         "XGBoost sur des données effectivement mesurées, évitant le biais circulaire "
         "des approches purement simulées."),
        ("Contribution 4 — Contexte africain",
         "Démonstration de la faisabilité d'un système de maintenance prédictive de la "
         "corrosion avec des composants disponibles sur le marché local camerounais "
         "(Jumia, marchés électroniques de Douala), ouvrant la voie à une "
         "démocratisation de la surveillance industrielle en Afrique."),
    ]

    for titre, desc in contribs:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(f"{titre} : ")
        r1.font.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        set_paragraph_spacing(p, before=0, after=6)

    add_heading(doc, "5.3  Extrapolation industrielle du prototype", 2)

    add_body(doc,
        "Le tableau d'analogie présenté à la section 3.1 permet de tracer directement "
        "le chemin de l'industrialisation du prototype. Chaque composant DIY peut être "
        "remplacé par son équivalent industriel certifié sans modifier l'architecture "
        "logicielle ni les algorithmes ML :"
    )

    extrap_data = [
        ("Fil de fer 0,2 mm",   "Sonde ER Cormon CW-20 (wire element)",   "Certification ATEX, durée de vie 5 ans"),
        ("ADS1115 + Wheatstone","Transmetteur Emerson Roxar 2600",          "Précision ±0,1 μm, sortie HART 4-20 mA"),
        ("ESP8266",             "Passerelle WirelessHART ou Modbus TCP",    "Protocoles industriels certifiés"),
        ("Supabase",            "PI Server OSIsoft/AVEVA",                  "Historien temps réel, conformité ISA-95"),
        ("Streamlit",           "SCADA Honeywell Experion ou Aspentech",    "Interface certifiée SIL2"),
        ("XGBoost Python",      "Module ML intégré dans SCADA ou MES",      "Déploiement en production"),
    ]

    table_ext = doc.add_table(rows=1, cols=3)
    table_ext.style = 'Table Grid'
    for i, txt in enumerate(["Composant prototype", "Équivalent industriel", "Apport de l'industrialisation"]):
        table_ext.rows[0].cells[i].text = txt
        table_ext.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table_ext.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
        table_ext.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    for row_d in extrap_data:
        row = table_ext.add_row()
        for i, txt in enumerate(row_d):
            row.cells[i].text = txt
            row.cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    add_body(doc,
        "Cette correspondance directe constitue l'argument central de la valeur académique "
        "de ce travail : le prototype n'est pas une démonstration isolée, mais un "
        "proof-of-concept industriellement transposable, dont chaque composant a été "
        "délibérément choisi en référence à son équivalent normalisé."
    )

    # Conclusion générale
    add_heading(doc, "CONCLUSION DE LA REVUE DE LITTÉRATURE", 1)

    add_body(doc,
        "Cette revue de littérature a établi le cadre théorique, normatif et technique "
        "du présent travail en suivant une progression rigoureuse du général au particulier. "
        "La corrosion, phénomène électrochimique universel quantifié par la loi de Faraday "
        "et encadré par un corpus normatif international (ASTM, NACE, API, ISO), représente "
        "un coût annuel de 2 500 milliards USD et demeure le principal mode de défaillance "
        "des infrastructures pétrolières."
    )

    add_body(doc,
        "Parmi les méthodes de surveillance disponibles, la sonde ER (ASTM G96) a été "
        "identifiée comme la seule technique combinant mesure continue, mesure directe du "
        "taux de corrosion et transposabilité en laboratoire à faible coût. Le prototype "
        "développé dans ce travail reproduit fidèlement le principe de la sonde ER "
        "industrielle avec des composants low-cost (fil de fer, pont de Wheatstone, "
        "ADS1115, ESP8266), dont la correspondance composant par composant avec les "
        "systèmes commerciaux a été établie."
    )

    add_body(doc,
        "L'apport original de ce travail réside dans l'intégration d'un modèle XGBoost "
        "à cette chaîne de mesure, transformant un outil de surveillance passif en un "
        "système de maintenance prédictive actif, capable de recommander automatiquement "
        "la dose d'inhibiteur optimale à partir des mesures en temps réel. Cette boucle "
        "fermée mesure → prédiction → recommandation constitue la contribution centrale "
        "de ce mémoire, en réponse à des gaps identifiés dans la littérature existante."
    )

    # ── SAUVEGARDE ─────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"OK  Document genere : {OUTPUT}")

if __name__ == "__main__":
    build()
