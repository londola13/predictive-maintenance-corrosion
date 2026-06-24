# -*- coding: utf-8 -*-
"""Génère le mémoire Word avec une mise en page soignée.

Pipeline :
  1. Fabrique un reference.docx (marges, police TNR 12 justifié 1,5, styles de titres,
     légendes, numéros de page) à partir du modèle pandoc.
  2. Prétraite memoire_v4.md (page de titre centrée, sauts de page, champ TOC Word).
  3. Lance pandoc -> Memoire_M2_Corrosion_BATOUMBI.docx

Usage : ../venv/Scripts/python.exe build_word.py
"""
import os
import re
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
os.chdir(HERE)
SRC = "memoire_v4.md"
TMP = "_memoire_tmp.md"
REF = "_reference.docx"
OUT = "Memoire_M2_Corrosion_BATOUMBI.docx"
NOIR = RGBColor(0x00, 0x00, 0x00)


# ───────────────────────── 1. REFERENCE.DOCX ─────────────────────────
def set_font(style, name="Times New Roman", size=None, bold=None, italic=None, color=NOIR):
    f = style.font
    f.name = name
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = color
    # forcer la police aussi pour les scripts complexes / east-asian
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def page_number_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for typ, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if typ:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), typ)
            run._r.append(fld)
        else:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            run._r.append(instr)


def build_reference():
    subprocess.run(f'pandoc --print-default-data-file reference.docx > "{REF}"',
                   shell=True, check=True)
    doc = Document(REF)

    # --- corps (Normal) : TNR 12, justifié, interligne 1,5 ---
    normal = doc.styles["Normal"]
    set_font(normal, size=12)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)

    # --- titres : noir, gras, hiérarchie, solidaires du paragraphe suivant ---
    for nom, size, italic, sb, sa in [
        ("Heading 1", 16, False, 18, 12),
        ("Heading 2", 14, False, 14, 8),
        ("Heading 3", 13, False, 10, 6),
        ("Heading 4", 12, True, 8, 4),
        ("Heading 5", 12, True, 6, 4),
    ]:
        try:
            s = doc.styles[nom]
        except KeyError:
            continue
        set_font(s, size=size, bold=True, italic=italic, color=NOIR)
        s.paragraph_format.space_before = Pt(sb)
        s.paragraph_format.space_after = Pt(sa)
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # --- légendes figures / tableaux : 10pt italique centré ---
    for nom in ("Image Caption", "Table Caption", "Caption"):
        try:
            s = doc.styles[nom]
        except KeyError:
            continue
        set_font(s, size=10, italic=True, color=NOIR)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_before = Pt(4)
        s.paragraph_format.space_after = Pt(10)

    # --- styles de table des matières (TOC 1..9) : ALIGNÉS À GAUCHE ---
    # Sinon, basés sur Normal (justifié), Word étire chaque entrée sur toute la largeur
    # (espacement disproportionné). Word apparie ces styles par styleId "TOC1".."TOC9".
    for i in range(1, 10):
        try:
            s = doc.styles.add_style(f"toc {i}", 1)
            s.element.set(qn("w:styleId"), f"TOC{i}")
            s.base_style = doc.styles["Normal"]
            set_font(s, size=12)
            s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            s.paragraph_format.space_after = Pt(2)
        except Exception:
            pass

    # --- style centré pour la page de titre ---
    try:
        pt_style = doc.styles.add_style("PageTitre", 1)  # 1 = paragraph
        pt_style.base_style = doc.styles["Normal"]
        set_font(pt_style, size=12)
        pt_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_style.paragraph_format.space_after = Pt(8)
    except Exception:
        pass

    # --- images centrées (le paragraphe qui contient l'image) ---
    try:
        body = doc.styles["Body Text"]
        set_font(body, size=12)
        body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        pass

    # --- mise en page : A4, marges (3 cm reliure à gauche), numéro de page ---
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    page_number_field(sec.footer.paragraphs[0])

    doc.save(REF)
    print("  reference.docx fabriqué")


# ───────────────────────── 2. PRÉTRAITEMENT ─────────────────────────
def convert_titre(src):
    """Convertit le tableau auteur de la page de titre (avant le 1er \\newpage)
    en lignes de texte simples (sinon il s'affiche en tableau bordé)."""
    parts = src.split("\\newpage", 1)
    titre, reste = parts[0], ("\\newpage" + parts[1] if len(parts) > 1 else "")
    out = []
    for line in titre.split("\n"):
        st = line.strip()
        if re.fullmatch(r"\|[\s|]*\|", st):                      # | | | (en-tête vide)
            continue
        if "-" in st and re.fullmatch(r"\|[-:\s|]+\|", st):      # |---|---| (séparateur)
            continue
        m = re.fullmatch(r"\|\s*(.+?)\s*\|\s*(.*?)\s*\|", st)    # | label | valeur |
        out.append(f"{m.group(1)} {m.group(2)}".rstrip() if m else line)
    return "\n".join(out) + reste


def preprocess():
    src = open(SRC, encoding="utf-8").read()

    # retirer le bloc YAML d'en-tête (sinon pandoc génère un "Title" + page parasite)
    src = re.sub(r"(?s)\A---\n.*?\n---\n", "", src)
    src = re.sub(r"(?s)\A\s*\\newpage[ \t]*\n", "", src)  # + le \newpage qui le suivait

    # commentaires décoratifs
    src = re.sub(r"(?m)^<!--.*-->[ \t]*$", "", src)

    # page de titre : retirer le titre "# PAGE DE TITRE", convertir le tableau auteur,
    # centrer tout le bloc jusqu'au premier \newpage
    src = re.sub(r"(?m)^# PAGE DE TITRE[ \t]*\n", "", src)
    src = convert_titre(src)
    # (centrage de la page de titre : fait en post-traitement sur le .docx)

    # LISTE DES FIGURES / TABLEAUX -> champs Word (table des illustrations par style)
    # NB: séparateur d'arguments ';' (locale FR de Word) — avec ',' Word FR lit "Style,1"
    # comme un seul nom de style et n'affiche "aucune entrée".
    def tof(style):
        return (
            "\n\n```{=openxml}\n"
            '<w:p><w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
            f'<w:r><w:instrText xml:space="preserve"> TOC \\h \\z \\t "{style};1" </w:instrText></w:r>'
            '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
            "<w:r><w:t>Liste générée par Word — sélectionner puis F9 pour mettre à jour.</w:t></w:r>"
            '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>\n```\n'
        )
    src = re.sub(r"(?m)^# LISTE DES FIGURES[ \t]*\n+\*\([^\n]*\)\*[ \t]*",
                 lambda m: "# LISTE DES FIGURES" + tof("Image Caption"), src)
    src = re.sub(r"(?m)^# LISTE DES TABLEAUX[ \t]*\n+\*\([^\n]*\)\*[ \t]*",
                 lambda m: "# LISTE DES TABLEAUX" + tof("Table Caption"), src)

    # champ TOC Word à l'emplacement "TABLE DES MATIÈRES"
    toc = (
        "\n\n```{=openxml}\n"
        '<w:p><w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        "<w:r><w:t>Clic droit puis « Mettre à jour les champs » (ou sélectionner + F9) "
        "pour générer la table.</w:t></w:r>"
        '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>\n'
        "```\n"
    )
    src = re.sub(r"(?m)^(# TABLE DES MATIÈRES[ \t]*)$", lambda m: m.group(1) + toc, src)

    # \newpage -> saut de page Word
    pb = ('\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n')
    src, n = re.subn(r"(?m)^\\newpage[ \t]*$", pb, src)

    open(TMP, "w", encoding="utf-8").write(src)
    print(f"  prétraitement OK ({n} sauts de page)")


# ───────────────────────── 3. PANDOC ─────────────────────────
def run_pandoc():
    cmd = [
        "pandoc", TMP, "-o", OUT,
        "--reference-doc", REF,
        "-f", "markdown+hard_line_breaks",
        "--resource-path", ".",
    ]
    subprocess.run(cmd, check=True)
    print(f"  pandoc -> {OUT}")


# ───────────────────────── 4. POST-TRAITEMENT (.docx) ─────────────────────────
def post_process():
    doc = Document(OUT)
    # 1) centrer la page de titre : tous les paragraphes jusqu'au 1er saut de page
    for p in doc.paragraphs:
        if p._p.xpath('.//w:br[@w:type="page"]'):
            break
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 2) centrer les figures (paragraphes contenant une image)
    for p in doc.paragraphs:
        if p._p.xpath(".//w:drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 3) labels de tableaux "Tableau X — ..." -> style "Table Caption"
    #    (pour alimenter la LISTE DES TABLEAUX) + centrés au-dessus du tableau
    for p in doc.paragraphs:
        if re.match(r"^Tableau\s+[\dIVXLA-Z]+", p.text.strip()):
            p.style = doc.styles["Table Caption"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 4) SOMMAIRE manuel (liste de sections) : aligner À GAUCHE — sinon justifié,
    #    chaque ligne est étirée sur toute la largeur (espacement disproportionné).
    in_somm = False
    for p in doc.paragraphs:
        nom = p.style.name if p.style else ""
        if nom.startswith("Heading") and p.text.strip() == "SOMMAIRE":
            in_somm = True
            continue
        if in_somm:
            if nom.startswith("Heading"):
                break
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if p._p.xpath('.//w:br[@w:type="page"]'):
                break
    doc.save(OUT)
    print("  post-traitement (page de titre + figures + SOMMAIRE)")


if __name__ == "__main__":
    print("Génération du mémoire Word…")
    build_reference()
    preprocess()
    run_pandoc()
    post_process()
    sz = os.path.getsize(OUT) / 1024
    print(f"\n>>> {OUT}  ({sz:.0f} Ko)")
